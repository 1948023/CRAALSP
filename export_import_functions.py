"""
Export and Import Functions for Risk Assessment Tool
Separated module for better code organization
"""

import csv
import json
import os
import sys
import re
import datetime
import logging
import traceback
from tkinter import messagebox, filedialog

def get_base_path():
    """Get the base path for the application (works with both .py and .exe)"""
    if getattr(sys, 'frozen', False):
        # Running as compiled executable
        # PyInstaller stores data files in sys._MEIPASS
        return os.path.dirname(sys.executable)
    else:
        # Running as script
        return os.path.dirname(os.path.abspath(__file__))

def get_output_path():
    """Get the path where output files should be saved"""
    if getattr(sys, 'frozen', False):
        # Running as compiled executable - save next to the .exe
        return os.path.dirname(sys.executable)
    else:
        # Running as script - save in script directory
        return os.path.dirname(os.path.abspath(__file__))

# Word export/import functionality
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ExportImportManager:
    """Handles all export and import operations"""
    
    def __init__(self, main_app):
        """Initialize with reference to main application"""
        self.app = main_app
    
    def export_csv(self):
        """Export analyzed threats to CSV files"""
        try:
            # Load threat details from Threat_Subset.csv
            threat_details = self._load_threat_details()
            
            # Get all analyzed threats
            analyzed_threats = self._get_analyzed_threats()
            
            if not analyzed_threats:
                messagebox.showinfo("Export", "No analyzed threats found. Please complete at least one threat assessment before exporting.")
                return
            
            # Create export folder with timestamp in Output directory
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = os.path.join(get_output_path(), "Output")
            os.makedirs(output_dir, exist_ok=True)
            export_folder = os.path.join(output_dir, f"CSV_Export_{timestamp}")
            os.makedirs(export_folder, exist_ok=True)
            
            # Export main threat file
            self._export_main_threat_file(analyzed_threats, threat_details, export_folder, timestamp)
            
            # Export asset-specific files
            created_files_count = self._export_asset_specific_files(analyzed_threats, threat_details, export_folder, timestamp)
            
            messagebox.showinfo("Export", f"Export completed successfully!\n\nFiles created in folder: {os.path.basename(export_folder)}\n- Threat_Analyzed.csv (main file)\n- {created_files_count} asset-specific files\n\nNote: Only assets with valid risk values were exported.")
            
        except Exception as e:
            messagebox.showerror("Export Error", f"An error occurred during export:\n{str(e)}")
            logging.error(f"Export error: {str(e)}")

    def export_word_report(self):
        """Export complete risk assessment report to Word document"""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Error", "python-docx library not available!\nInstall with: pip install python-docx")
            return
            
        try:
            # Get analyzed threats
            analyzed_threats = self._get_analyzed_threats()
            
            if not analyzed_threats:
                messagebox.showinfo("Export", "No analyzed threats found. Complete at least one assessment before exporting.")
                return
            
            # Create filename with timestamp
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Risk_Assessment_Report_{timestamp}.docx"
            
            # Create Output directory if it doesn't exist
            output_dir = os.path.join(get_output_path(), "Output")
            os.makedirs(output_dir, exist_ok=True)
            
            file_path = os.path.join(output_dir, filename)
            
            # Create Word document
            doc = Document()
            
            # Set narrow margins for the document
            from docx.shared import Inches
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)      # 1.27 cm
                section.bottom_margin = Inches(0.5)   # 1.27 cm
                section.left_margin = Inches(0.5)     # 1.27 cm
                section.right_margin = Inches(0.5)    # 1.27 cm
            
            # Add content sections
            self._add_word_title_and_info(doc)
            self._add_main_threats_overview_table(doc, analyzed_threats)
            self._add_asset_assessment_table(doc, analyzed_threats)
            self._add_detailed_threat_analysis(doc, analyzed_threats)
            
            # Save document
            doc.save(file_path)
            
            messagebox.showinfo("Success", f"Risk Assessment Report exported to:\n{file_path}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error during Word export:\n{str(e)}")
            logging.error(f"Word export error: {str(e)}")

    def import_word_report(self):
        """Import risk assessment data from Word document"""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Error", "python-docx library not available!\nInstall with: pip install python-docx")
            return
            
        try:
            # Select file to import
            file_path = filedialog.askopenfilename(
                title="Select Report to Import",
                filetypes=[("Word Documents", "*.docx"), ("All files", "*.*")],
                initialdir=get_base_path()
            )
            
            if not file_path:
                return
            
            # Clear existing data
            self.app.threat_data = {}
            self.app.asset_data = {}
            
            # Load and parse document
            doc = Document(file_path)
            self._parse_word_document_simple(doc)
            
            # Update main table
            self.app.update_all_threats_in_main_table()
            
            # Show summary
            threat_count = len(self.app.threat_data)
            asset_count = 0
            
            # Count actual imported asset assessments
            if hasattr(self.app, 'asset_data') and self.app.asset_data:
                for timestamp_key, assessment_data in self.app.asset_data.items():
                    if timestamp_key.startswith('imported_'):
                        asset_count = len([k for k in assessment_data.keys() if k.endswith('_probability')])
                        break  # Only count the latest import
            
            # Count threat-asset combinations
            threat_asset_combinations = 0
            for threat_name, threat_data in self.app.threat_data.items():
                threat_asset_combinations += len(threat_data)
            
            success_msg = f"Report imported successfully from: {os.path.basename(file_path)}\n\n"
            success_msg += f"Data imported:\n"
            success_msg += f"- {threat_count} unique threats\n"
            success_msg += f"- {threat_asset_combinations} threat-asset assessments\n"
            success_msg += f"- {asset_count} asset assessments\n\n"
            success_msg += f"All detailed criteria have been imported and are ready for analysis."

            messagebox.showinfo("Success", success_msg)
            
            logging.info(f"Import completed: {threat_count} threats, {asset_count} asset assessments")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error during import:\n{str(e)}")
            logging.error(f"Word import error: {str(e)}")

    def legacy_import(self):
        """Import legacy Word reports from Risk_Assessment_II_Phase_Optimized.py and map old threats to new ones"""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Error", "python-docx library not available!\nInstall with: pip install python-docx")
            return
            
        try:
            # Select legacy report file to import
            file_path = filedialog.askopenfilename(
                title="Select Legacy Report to Import",
                filetypes=[("Word Documents", "*.docx"), ("All files", "*.*")],
                initialdir=get_base_path()
            )
            
            if not file_path:
                return
            
            # Load threat mapping from Legacy.csv
            threat_mapping = self._load_threat_mapping()
            if not threat_mapping:
                messagebox.showerror("Error", "Could not load threat mapping from Legacy.csv")
                return
            
            # Setup logging with UTF-8 encoding
            logging.basicConfig(
                level=logging.INFO,
                format='%(levelname)s: %(message)s'
            )
            self.logger = logging.getLogger(__name__)
            
            # Load and parse legacy document
            doc = Document(file_path)
            legacy_data = self._parse_legacy_document(doc, threat_mapping)
            
            if not legacy_data:
                messagebox.showwarning("Warning", "No valid legacy data found in the document")
                return
            
            # Import the mapped data into current threat_data structure
            imported_count = self._integrate_legacy_data(legacy_data)
            
            # Show success message
            success_msg = f"Legacy report imported successfully!\n\n"
            success_msg += f"File: {os.path.basename(file_path)}\n"
            success_msg += f"Imported {imported_count} threat-asset assessments\n"
            success_msg += f"Mapped to {len(legacy_data)} current threats\n\n"
            success_msg += "Data has been integrated into the current threat analysis."
            
            messagebox.showinfo("Success", success_msg)
            self.logger.info(f"Legacy import completed: {imported_count} assessments imported")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error during legacy import:\n{str(e)}")
            logging.error(f"Legacy import error: {str(e)}")
            traceback.print_exc()

    def _normalize_threat_name(self, name):
        """Normalize threat names for consistent mapping"""
        if not name:
            return ""
        
        # Remove extra spaces, convert to lowercase, remove special characters
        normalized = name.strip().lower()
        # Remove common separators and spaces
        normalized = normalized.replace(' ', '').replace('/', '').replace('-', '').replace('_', '')
        # Remove parentheses and their contents
        import re
        normalized = re.sub(r'\([^)]*\)', '', normalized)
        
        return normalized


    def _load_threat_mapping(self):
        """Load threat mapping from Legacy.csv"""
        threat_mapping = {}
        legacy_file = os.path.join(get_base_path(), "Legacy.csv")
        
        try:
            with open(legacy_file, 'r', newline='', encoding='utf-8') as csvfile:
                # Read the first line to check the format
                first_line = csvfile.readline().strip()
                csvfile.seek(0)  # Reset to beginning
                
                # Check if it uses '; ' (semicolon with space) as separator
                if '; ' in first_line:
                    # Manual parsing for '; ' separator
                    lines = csvfile.readlines()
                    headers = [h.strip() for h in lines[0].strip().split(';')]
                    
                    for line_num, line in enumerate(lines[1:], 2):
                        if line.strip():
                            parts = [p.strip() for p in line.strip().split(';')]
                            if len(parts) >= 2:
                                old_threat = parts[0].strip()
                                new_threat = parts[1].strip()
                                
                                if old_threat and new_threat:
                                    # Normalize the old threat name for consistent mapping
                                    normalized_old = self._normalize_threat_name(old_threat)
                                    if normalized_old not in threat_mapping:
                                        threat_mapping[normalized_old] = []
                                    threat_mapping[normalized_old].append(new_threat)
                                    logging.info(f"Mapped: '{old_threat}' (normalized: '{normalized_old}') -> '{new_threat}'")
                else:
                    # Standard CSV parsing
                    reader = csv.DictReader(csvfile, delimiter=';')
                    for row in reader:
                        old_threat = row.get('Old Threat', '').strip()
                        new_threat = row.get('New Threat', '').strip()
                        
                        if old_threat and new_threat:
                            # Normalize the old threat name for consistent mapping
                            normalized_old = self._normalize_threat_name(old_threat)
                            if normalized_old not in threat_mapping:
                                threat_mapping[normalized_old] = []
                            threat_mapping[normalized_old].append(new_threat)
                            logging.info(f"Mapped: '{old_threat}' (normalized: '{normalized_old}') -> '{new_threat}'")
            
            logging.info(f"Loaded {len(threat_mapping)} threat mappings from Legacy.csv")
            
            # Debug: print all loaded mappings
            for normalized_key, new_threats in threat_mapping.items():
                logging.info(f"  '{normalized_key}' -> {new_threats}")
                
            return threat_mapping
            
        except FileNotFoundError:
            logging.error(f"Legacy mapping file not found: {legacy_file}")
            return {}
        except Exception as e:
            logging.error(f"Error loading threat mapping: {e}")
            return {}

    def _parse_legacy_document(self, doc, threat_mapping):
        """Parse legacy Word document generated by Risk_Assessment_II_Phase_Optimized.py"""
        try:
            # Setup logging
            if not hasattr(self, 'logger'):
                self.logger = logging.getLogger(__name__)
            
            legacy_data = {}
            
            # Parse document structure in order
            all_elements = self._extract_document_elements(doc)
            
            # Process elements to find threat tables
            in_detailed_section = False
            current_threat = None
            
            for element_type, element_data in all_elements:
                if element_type == 'paragraph':
                    text = element_data.strip()
                    
                    # Check for "Detailed Threat Analysis" section
                    if "Detailed Threat Analysis" in text:
                        in_detailed_section = True
                        self.logger.info("[OK] Found Detailed Threat Analysis section")
                        continue
                    
                    # If in detailed section, check for threat names (heading level 2)
                    if in_detailed_section and text:
                        # Check if text matches any threat name in mapping (normalized comparison)
                        normalized_text = self._normalize_threat_name(text)
                        if normalized_text in threat_mapping:
                            current_threat = text  # Keep original name for logging
                            self.logger.info(f"📋 Found legacy threat: {current_threat} (normalized: {normalized_text})")
                            continue
                        
                elif element_type == 'table' and current_threat and in_detailed_section:
                    table = element_data
                    
                    # Check if this is an asset assessment table (9 columns)
                    if len(table.columns) == 9:
                        self.logger.info(f"🔍 Processing asset table for threat: {current_threat}")
                        
                        # Parse the table data
                        table_data = self._parse_legacy_asset_table(table, current_threat)
                        
                        # Map legacy threat to new threats using normalized name
                        normalized_threat = self._normalize_threat_name(current_threat)
                        new_threat_names = threat_mapping.get(normalized_threat, [])
                        
                        if new_threat_names and table_data:
                            # Apply to EACH mapped threat separately (not all at once)
                            for new_threat in new_threat_names:
                                if new_threat not in legacy_data:
                                    legacy_data[new_threat] = {}
                                
                                # Create a copy of table_data for this specific threat
                                threat_specific_data = {}
                                for asset_key, criteria in table_data.items():
                                    threat_specific_data[asset_key] = criteria.copy()
                                
                                legacy_data[new_threat].update(threat_specific_data)
                                self.logger.info(f"[OK] Mapped '{current_threat}' -> '{new_threat}' with {len(threat_specific_data)} assets")
                        else:
                            self.logger.warning(f"[ERROR] No mapping found for threat: {current_threat} (normalized: {normalized_threat})")
                    
                    elif len(table.columns) == 2:
                        # Security controls table - skip
                        self.logger.info(f"   → Skipping security controls table")
                    
                    else:
                        self.logger.warning(f"   → Unknown table format ({len(table.columns)} columns)")
            
            self.logger.info(f"🎯 Legacy parsing completed. Found data for {len(legacy_data)} threats")
            return legacy_data
            
        except Exception as e:
            self.logger.error(f"[ERROR] Error parsing legacy document: {str(e)}")
            return {}

    def _extract_document_elements(self, doc):
        """Extract all document elements (paragraphs and tables) in order"""
        all_elements = []
        
        try:
            # Process document body elements in order
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    # Find corresponding paragraph object
                    for para in doc.paragraphs:
                        if para._element == element:
                            all_elements.append(('paragraph', para.text))
                            break
                            
                elif element.tag.endswith('tbl'):  # Table
                    # Find corresponding table object
                    for table in doc.tables:
                        if table._element == element:
                            all_elements.append(('table', table))
                            break
                            
        except Exception as e:
            self.logger.error(f"Error extracting document elements: {str(e)}")
            # Fallback: just process tables
            for table in doc.tables:
                all_elements.append(('table', table))
        
        return all_elements

    def _parse_legacy_asset_table(self, table, threat_name):
        """Parse legacy asset assessment table with 9 columns"""
        try:
            asset_data = {}
            
            # Verify table structure
            if len(table.rows) < 2:
                self.logger.warning(f"Table too short for threat {threat_name}")
                return asset_data
                
            # Expected headers: Asset, Vulnerability, Access Control, Defense Capability, 
            #                  Operational Impact, Recovery Time, Likelihood, Impact, Risk Level
            expected_headers = ['asset', 'vulnerability', 'access', 'defense', 'operational', 'recovery']
            
            # Get actual headers
            header_row = table.rows[0]
            actual_headers = [cell.text.strip().lower() for cell in header_row.cells]
            
            # Verify this is the right table format
            header_match_count = 0
            for expected in expected_headers:
                if any(expected in header for header in actual_headers):
                    header_match_count += 1
            
            if header_match_count < 4:
                self.logger.warning(f"Table header mismatch for threat {threat_name}: {actual_headers}")
                return asset_data
            
            # Process each data row
            for row_idx in range(1, len(table.rows)):
                row = table.rows[row_idx]
                cells = [cell.text.strip() for cell in row.cells]
                
                if len(cells) < 6:
                    continue
                
                # Extract asset name (which is actually a sub-category in legacy system)
                legacy_asset_name = cells[0]
                if not legacy_asset_name or legacy_asset_name in ["", "N/A", "No asset data available"]:
                    continue
                
                # Extract criteria scores (columns 1-5)
                criteria_scores = {}
                valid_criteria = 0
                
                # Mapping personalizzato
                mapping = {
                    0: ["0"],        # Vulnerability → Vulnerability (criterio 0)
                    1: ["3", "4"],   # Access → Access (criterio 3) + Privilege (criterio 4)  
                    2: ["1", "2"],   # Defense → Mitigation (criterio 1) + Detection (criterio 2)
                    3: ["5"],        # Operational → Response (criterio 5)
                    4: ["6"]         # Recovery → Resilience (criterio 6)
                }

                for i in range(1, 6):
                    if i < len(cells):
                        score = self._parse_score_from_cell(cells[i])
                        if score is not None:
                            # Applica il punteggio a tutti i criteri mappati
                            target_criteria = mapping.get(i-1, [])
                            for target_criterion in target_criteria:
                                criteria_scores[target_criterion] = str(score)
                            valid_criteria += len(target_criteria)

                
                # Only proceed if we have at least 3 valid criteria
                if valid_criteria < 3:
                    self.logger.warning(f"   Insufficient criteria for legacy asset '{legacy_asset_name}' ({valid_criteria}/5)")
                    continue
                
                # Find all assets in current system that match this legacy sub-category
                matching_assets = self._find_assets_by_subcategory(legacy_asset_name)
                
                if not matching_assets:
                    self.logger.warning(f"No matching assets found for legacy asset '{legacy_asset_name}'")
                    continue
                
                # Apply criteria to all matching assets
                for asset_index in matching_assets:
                    asset_key = f"{asset_index + 1}_probability"
                    asset_data[asset_key] = criteria_scores.copy()
                    
                    # Get the actual asset name for logging
                    if asset_index < len(self._get_current_asset_categories()):
                        _, _, asset_name = self._get_current_asset_categories()[asset_index]
                        self.logger.info(f"   Applied legacy asset '{legacy_asset_name}' criteria to '{asset_name}' (index {asset_index})")
            
            self.logger.info(f"   Total assets configured: {len(asset_data)}")
            return asset_data
            
        except Exception as e:
            self.logger.error(f"Error parsing asset table for threat {threat_name}: {str(e)}")
            return {}

    def _find_assets_by_subcategory(self, legacy_asset_name):
        """Find all asset indices that match the legacy asset name (sub-category)"""
        matching_indices = []
        
        # Get current asset categories
        current_assets = self._get_current_asset_categories()
        
        # Try exact match with sub-category first
        for i, (category, sub_category, component) in enumerate(current_assets):
            if sub_category and sub_category.lower() == legacy_asset_name.lower():
                matching_indices.append(i)
        
        # If no exact match, try partial matches
        if not matching_indices:
            for i, (category, sub_category, component) in enumerate(current_assets):
                # Check if legacy name is in sub-category or vice versa
                if sub_category and (
                    legacy_asset_name.lower() in sub_category.lower() or 
                    sub_category.lower() in legacy_asset_name.lower()
                ):
                    matching_indices.append(i)
        
        # Special case mappings for known legacy sub-categories
        if not matching_indices:
            legacy_mappings = {
                "ground stations": ["ground stations"],
                "mission control": ["mission control"],
                "data processing centers": ["data processing centers"],
                "remote terminals": ["remote terminals"],
                "user ground segment": ["user ground segment"],
                "platform": ["platform"],
                "payload": ["payload"],
                "link": ["link"],
                "user": ["user"]
            }
            
            legacy_lower = legacy_asset_name.lower()
            for mapping_key, target_subcategories in legacy_mappings.items():
                if mapping_key in legacy_lower:
                    for i, (category, sub_category, component) in enumerate(current_assets):
                        if sub_category and any(target in sub_category.lower() for target in target_subcategories):
                            matching_indices.append(i)
                    break
        
        return matching_indices

    def _get_current_asset_categories(self):
        """Get current asset categories from the application"""
        if hasattr(self.app, 'ASSET_CATEGORIES'):
            return self.app.ASSET_CATEGORIES
        
        # Fallback: standard asset categories
        return [
            ("Ground", "Ground Stations", "Tracking"), ("Ground", "Ground Stations", "Ranging"),
            ("Ground", "Mission Control", "Telemetry processing"), ("Ground", "Mission Control", "Commanding"),
            ("Ground", "Data Processing Centers", "Mission Analysis"), ("Ground", "Remote Terminals", "Network access"),
            ("Ground", "User Ground Segment", "Development"), ("Space", "Platform", "Bus"),
            ("Space", "Payload", "Instruments"), ("Link", "", "Uplink"), ("Link", "", "Downlink"),
            ("User", "", "End User")
        ]

    def _parse_score_from_cell(self, cell_text):
        """Parse score from legacy table cell - handles 'Score X' format"""
        if not cell_text:
            return None
        
        text = cell_text.strip()
        
        # Handle "Score X" format
        if text.lower().startswith("score"):
            try:
                score_part = text.lower().replace("score", "").strip()
                score = int(score_part)
                if 1 <= score <= 5:
                    return score
            except ValueError:
                pass
        
        # Handle direct number
        if text.isdigit():
            score = int(text)
            if 1 <= score <= 5:
                return score
        
        # Handle "N/A" or empty
        if text.lower() in ['n/a', 'na', '-', '']:
            return None
        
        # Try to extract number from text
        import re
        numbers = re.findall(r'\b([1-5])\b', text)
        if numbers:
            return int(numbers[0])
        
        return None

    def _integrate_legacy_data(self, legacy_data):
        """Integrate legacy data into current threat analysis"""
        try:
            imported_count = 0
            
            # Get current system threat names
            current_threats = self._get_current_threat_names()
            
            for new_threat_name, asset_assessments in legacy_data.items():
                # Check if threat exists in current system
                if new_threat_name not in current_threats:
                    self.logger.warning(f"Threat '{new_threat_name}' not found in current system")
                    continue
                
                # Initialize threat data if needed
                if new_threat_name not in self.app.threat_data:
                    self.app.threat_data[new_threat_name] = {}
                
                # Merge asset assessments
                for asset_key, assessment_data in asset_assessments.items():
                    self.app.threat_data[new_threat_name][asset_key] = assessment_data
                    imported_count += 1
            
            # Update the main table display
            self.app.update_all_threats_in_main_table()
            
            self.logger.info(f"[OK] Integrated {imported_count} assessments into current system")
            return imported_count
            
        except Exception as e:
            self.logger.error(f"[ERROR] Error integrating legacy data: {str(e)}")
            return 0

    def _get_current_threat_names(self):
        """Get list of threat names from current system"""
        try:
            # Load current threats from Threat.csv
            current_threats = set()
            threats_file = os.path.join(get_base_path(), "Threat.csv")
            
            with open(threats_file, 'r', newline='', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile, delimiter=';')
                for row in reader:
                    threat_name = row.get('THREAT', '').strip()
                    if threat_name:
                        current_threats.add(threat_name)
            
            return list(current_threats)
            
        except Exception as e:
            self.logger.error(f"Error loading current threat names: {str(e)}")
            return []

    # ===== PRIVATE HELPER METHODS =====
    
    def _load_threat_details(self):
        """Load threat details from Threat_Subset.csv"""
        threat_details = {}
        threats_file = os.path.join(get_base_path(), "Threat_Subset.csv")
        
        try:
            with open(threats_file, 'r', newline='', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile, delimiter=';')
                for row in reader:
                    threat_name = row.get('THREAT', '').strip()
                    if threat_name:
                        threat_details[threat_name] = {
                            'category': row.get('THREAT CATEGORY', '').strip(),
                            'description': row.get('THREAT DESCRIPTION', '').strip()
                        }
        except FileNotFoundError:
            logging.warning(f"Threat details file not found: {threats_file}")
        except Exception as e:
            logging.error(f"Error loading threat details: {e}")
        
        return threat_details

    def _get_analyzed_threats(self):
        """Get list of threats that have been analyzed"""
        analyzed_threats = []
        
        for threat_name in sorted(self.app.threat_data.keys()):
            threat_data = self.app.threat_data[threat_name]
            has_valid_risk = False
            
            for asset_key, asset_data in threat_data.items():
                likelihood = self.app.calculate_likelihood_from_saved_data(asset_data)
                impact = self.app.calculate_impact_from_saved_data(asset_data)
                
                if likelihood >= 0 and impact >= 0:
                    has_valid_risk = True
                    break
            
            if has_valid_risk:
                analyzed_threats.append(threat_name)
        
        return analyzed_threats

    def _export_main_threat_file(self, analyzed_threats, threat_details, export_folder, timestamp):
        """Export main Threat_Analyzed.csv file"""
        filename = os.path.join(export_folder, f"Threat_Analyzed.csv")
        
        with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile, delimiter=';')
            writer.writerow(["THREAT", "Likelihood", "Impact", "Risk"])
            
            for threat_name in analyzed_threats:
                max_likelihood, max_impact, max_risk = self.app.get_threat_max_risk(threat_name)
                
                if max_risk and max_risk != "":
                    writer.writerow([threat_name, max_likelihood, max_impact, max_risk])

    def _export_asset_specific_files(self, analyzed_threats, threat_details, export_folder, timestamp):
        """Export asset-specific CSV files"""
        analyzed_assets = self._get_analyzed_assets()
        created_files = 0
        
        for asset_name in analyzed_assets:
            asset_threats_data = []
            
            for threat_name in analyzed_threats:
                asset_likelihood, asset_impact, asset_risk = self.app.get_threat_asset_risk(threat_name, asset_name)
                
                if asset_likelihood and asset_impact and asset_risk and asset_risk != "":
                    asset_threats_data.append({
                        'threat': threat_name,
                        'likelihood': asset_likelihood,
                        'impact': asset_impact,
                        'risk': asset_risk
                    })
            
            if asset_threats_data:
                filename = os.path.join(export_folder, 
                                      f"Threat_Analyzed_{asset_name.replace('/', '_').replace(' ', '_')}.csv")
                
                with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
                    writer = csv.writer(csvfile, delimiter=';')
                    writer.writerow(["THREAT", "Likelihood", "Impact", "Risk"])
                    
                    for threat_data in asset_threats_data:
                        writer.writerow([
                            threat_data['threat'],
                            threat_data['likelihood'],
                            threat_data['impact'],
                            threat_data['risk']
                        ])
                
                created_files += 1
        
        return created_files

    def _get_analyzed_assets(self):
        """Get list of assets that have been analyzed"""
        analyzed_assets = set()
        
        for threat_name in self.app.threat_data.keys():
            threat_data = self.app.threat_data[threat_name]
            
            for asset_key, asset_data in threat_data.items():
                likelihood = self.app.calculate_likelihood_from_saved_data(asset_data)
                impact = self.app.calculate_impact_from_saved_data(asset_data)
                
                if likelihood >= 0 and impact >= 0:
                    asset_index = int(asset_key.split('_')[0]) - 1
                    if 0 <= asset_index < len(self.app.ASSET_CATEGORIES):
                        asset_name = self.app.ASSET_CATEGORIES[asset_index][2]
                        analyzed_assets.add(asset_name)
        
        return list(analyzed_assets)
    
    def _get_all_assets_with_likelihood_impact(self):
        """Get all assets that have both likelihood and impact values defined"""
        assets_with_assessment = set()  # Use set to avoid duplicates
        for i, (category, sub_category, asset_name) in enumerate(self.app.ASSET_CATEGORIES):
            asset_likelihood, asset_impact = self._get_asset_likelihood_impact(asset_name)
            if asset_likelihood and asset_impact:
                assets_with_assessment.add(asset_name)
        return sorted(list(assets_with_assessment))

    def _get_asset_detailed_criteria(self, asset_name):
        """Get detailed criteria scores for an asset from latest assessment"""
        if not self.app.asset_data:
            return {}
        
        # Find the most recent assessment - prioritize assessment_ keys over imported_ keys
        assessment_keys = [key for key in self.app.asset_data.keys() if key.startswith('assessment_')]
        imported_keys = [key for key in self.app.asset_data.keys() if key.startswith('imported_')]
        
        # Use the latest assessment key if available, otherwise use latest imported key
        if assessment_keys:
            latest_key = max(assessment_keys)
        elif imported_keys:
            latest_key = max(imported_keys)
        else:
            latest_key = max(self.app.asset_data.keys()) if self.app.asset_data else None
        
        if not latest_key or latest_key not in self.app.asset_data:
            return {}
        
        # Find asset index by name
        asset_index = -1
        for i, (category, sub_category, component) in enumerate(self.app.ASSET_CATEGORIES):
            if component == asset_name:
                asset_index = i
                break
        
        if asset_index == -1:
            return {}
        
        asset_key = f"{asset_index + 1}_probability"
        asset_assessment = self.app.asset_data[latest_key]
        
        if asset_key not in asset_assessment:
            return {}
        
        return asset_assessment[asset_key]

    def _get_threat_detailed_criteria(self, threat_name, asset_name):
        """Get detailed criteria scores for a threat-asset combination"""
        if threat_name not in self.app.threat_data:
            return {}
        
        # Find asset index by name
        asset_index = -1
        for i, (category, sub_category, component) in enumerate(self.app.ASSET_CATEGORIES):
            if component == asset_name:
                asset_index = i
                break
        
        if asset_index == -1:
            return {}
        
        asset_key = f"{asset_index + 1}_probability"
        threat_data = self.app.threat_data[threat_name]
        
        if asset_key not in threat_data:
            return {}
        
        return threat_data[asset_key]

    def _add_word_title_and_info(self, doc):
        """Add title and info to Word document"""
        title = doc.add_heading('Risk Assessment', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(f'Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()

    def _add_main_threats_overview_table(self, doc, analyzed_threats):
        """Add main threats overview table to Word document"""
        doc.add_heading('Main Threats Overview', level=1)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Threat'
        header_cells[1].text = 'Likelihood'
        header_cells[2].text = 'Impact'
        header_cells[3].text = 'Risk Level'
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Data
        for threat_name in analyzed_threats:
            max_likelihood, max_impact, max_risk = self.app.get_threat_max_risk(threat_name)
            
            row_cells = table.add_row().cells
            row_cells[0].text = threat_name
            row_cells[1].text = max_likelihood if max_likelihood else ""
            row_cells[2].text = max_impact if max_impact else ""
            row_cells[3].text = max_risk if max_risk else ""
        
        doc.add_paragraph()

    def _add_asset_assessment_table(self, doc, analyzed_threats):
        """Add asset assessment table to Word document with detailed criteria"""
        doc.add_heading('Asset Assessment Overview', level=1)
        
        # Get ALL assets with valid likelihood and impact (not just those in threat assessment)
        assets_with_assessment = self._get_all_assets_with_likelihood_impact()
        
        if not assets_with_assessment:
            doc.add_paragraph("No asset assessment data available.")
            return
        
        # Asset criteria names (9 criteria: 4 for likelihood, 5 for impact)
        asset_criteria = ['Dependency', 'Penetration', 'Maturity', 'Trust', 
                         'Performance', 'Schedule', 'Costs', 'Reputation', 'Recovery']
        
        # Create table with 15 columns: Category, Sub-category, Asset + 9 criteria + Likelihood, Impact, Risk
        table = doc.add_table(rows=1, cols=15)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        header_cells = table.rows[0].cells
        headers = ['Category', 'Sub-category', 'Asset'] + asset_criteria + ['Likelihood', 'Impact', 'Risk']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Data
        for asset_name in assets_with_assessment:
            asset_category, asset_sub_category = "", ""
            for category, sub_category, component in self.app.ASSET_CATEGORIES:
                if component == asset_name:
                    asset_category = category
                    asset_sub_category = sub_category
                    break
            
            # Get asset criteria details
            asset_criteria_data = self._get_asset_detailed_criteria(asset_name)
            asset_likelihood, asset_impact = self._get_asset_likelihood_impact(asset_name)
            asset_risk = ""
            if asset_likelihood and asset_impact:
                asset_risk = self.app.RISK_MATRIX.get((asset_likelihood, asset_impact), "")

            row_cells = table.add_row().cells
            row_cells[0].text = asset_category
            row_cells[1].text = asset_sub_category
            row_cells[2].text = asset_name
            
            # Add criteria scores (columns 3-11)
            for i in range(9):
                criteria_value = asset_criteria_data.get(str(i), "")
                row_cells[3 + i].text = criteria_value
            
            # Add likelihood, impact, risk (columns 12-14)
            row_cells[12].text = asset_likelihood if asset_likelihood else ""
            row_cells[13].text = asset_impact if asset_impact else ""
            row_cells[14].text = asset_risk if asset_risk else ""
        
        doc.add_paragraph()

    def _add_detailed_threat_analysis(self, doc, analyzed_threats):
        """Add detailed threat analysis to Word document"""
        doc.add_heading('Detailed Threat Analysis', level=1)
        
        for threat_name in analyzed_threats:
            doc.add_heading(f'{threat_name}', level=2)
            self._add_threat_risk_assessment_table(doc, threat_name)
            self._add_threat_controls_table(doc, threat_name)
            doc.add_paragraph()

    def _add_threat_risk_assessment_table(self, doc, threat_name):
        """Add risk assessment table for specific threat with detailed criteria"""
        doc.add_heading(f'Risk Assessment for {threat_name}', level=3)
        
        # Asset criteria (9) + Threat criteria (7) + summary columns
        asset_criteria = ['Dependency', 'Penetration', 'Maturity', 'Trust', 
                         'Performance', 'Schedule', 'Costs', 'Reputation', 'Recovery']
        threat_criteria = ['Vulnerability', 'Mitigation', 'Detection', 'Access', 
                          'Privilege', 'Response', 'Resilience']
        
        # Create table: Asset + 9 Asset criteria + Asset L/I + 7 Threat criteria + Threat L/I/R = 22 columns
        table = doc.add_table(rows=1, cols=22)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        header_cells = table.rows[0].cells
        headers = (['Asset'] + 
                  [f'A-{crit}' for crit in asset_criteria] + 
                  ['Asset Likelihood', 'Asset Impact'] +
                  [f'T-{crit}' for crit in threat_criteria] + 
                  ['Threat Likelihood', 'Threat Impact', 'Threat Risk Level'])
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Data
        assets_added = 0
        if threat_name in sorted(self.app.threat_data):
            threat_data = self.app.threat_data[threat_name]
            
            for asset_key, threat_asset_data in threat_data.items():
                try:
                    asset_index = int(asset_key.split('_')[0]) - 1
                    if 0 <= asset_index < len(self.app.ASSET_CATEGORIES):
                        category, sub_category, asset_name = self.app.ASSET_CATEGORIES[asset_index]
                        
                        # Threat likelihood and impact
                        threat_likelihood = self.app.calculate_likelihood_from_saved_data(threat_asset_data)
                        threat_impact = self.app.calculate_impact_from_saved_data(threat_asset_data)
                        
                        if threat_likelihood >= 0 and threat_impact >= 0:
                            threat_likelihood_cat = self.app.value_to_category(threat_likelihood)
                            threat_impact_cat = self.app.value_to_category(threat_impact)
                            threat_risk_level = self.app.RISK_MATRIX.get((threat_likelihood_cat, threat_impact_cat), "")
                            
                            # Asset likelihood and impact
                            asset_likelihood, asset_impact = self._get_asset_likelihood_impact(asset_name)
                            
                            # Get detailed criteria
                            asset_criteria_data = self._get_asset_detailed_criteria(asset_name)
                            threat_criteria_data = self._get_threat_detailed_criteria(threat_name, asset_name)
                            
                            row_cells = table.add_row().cells
                            col_idx = 0
                            
                            # Asset name
                            row_cells[col_idx].text = asset_name
                            col_idx += 1
                            
                            # Asset criteria (9 columns)
                            for i in range(9):
                                criteria_value = asset_criteria_data.get(str(i), "")
                                row_cells[col_idx].text = criteria_value
                                col_idx += 1
                            
                            # Asset likelihood and impact
                            row_cells[col_idx].text = asset_likelihood if asset_likelihood else "N/A"
                            col_idx += 1
                            row_cells[col_idx].text = asset_impact if asset_impact else "N/A"
                            col_idx += 1
                            
                            # Threat criteria (7 columns)
                            for i in range(7):
                                criteria_value = threat_criteria_data.get(str(i), "")
                                row_cells[col_idx].text = criteria_value
                                col_idx += 1
                            
                            # Threat likelihood, impact, risk
                            row_cells[col_idx].text = threat_likelihood_cat
                            col_idx += 1
                            row_cells[col_idx].text = threat_impact_cat
                            col_idx += 1
                            row_cells[col_idx].text = threat_risk_level
                            
                            assets_added += 1
                            
                except (ValueError, IndexError):
                    continue
        
        if assets_added == 0:
            row_cells = table.add_row().cells
            row_cells[0].text = "No risk assessment data available"
            for i in range(1, 22):
                row_cells[i].text = "N/A"
        
        doc.add_paragraph()

    def _add_threat_controls_table(self, doc, threat_name):
        """Add mitigation controls table for specific threat"""
        # Load controls for this threat
        controls = self.app.load_controls_for_threat(threat_name)
        
        if not controls:
            doc.add_paragraph(f"No specific mitigation controls found for {threat_name}.")
            return
        
        doc.add_heading(f'Mitigation Controls for {threat_name}', level=3)
        
        # Create controls table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        header_cells = table.rows[0].cells
        headers = ['Control Title', 'Control ID', 'Description', 'Reference Frameworks', 'Lifecycle Phase', 'Segment']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Data rows
        for control in controls:
            row_cells = table.add_row().cells
            row_cells[0].text = control.get('title', '')
            row_cells[1].text = control.get('control', '')
            row_cells[2].text = control.get('description', '')
            row_cells[3].text = control.get('reference', '')
            row_cells[4].text = control.get('lifecycle', '')
            row_cells[5].text = control.get('segment', '')
        
        doc.add_paragraph()

    def _get_asset_likelihood_impact(self, asset_name):
        """Get asset likelihood and impact from latest assessment"""
        if not self.app.asset_data:
            return "", ""
        
        # Find the most recent assessment - prioritize assessment_ keys over imported_ keys
        assessment_keys = [key for key in self.app.asset_data.keys() if key.startswith('assessment_')]
        imported_keys = [key for key in self.app.asset_data.keys() if key.startswith('imported_')]
        
        # Use the latest assessment key if available, otherwise use latest imported key
        if assessment_keys:
            latest_key = max(assessment_keys)
        elif imported_keys:
            latest_key = max(imported_keys)
        else:
            latest_key = max(self.app.asset_data.keys()) if self.app.asset_data else None
        
        if not latest_key or latest_key not in self.app.asset_data:
            return "", ""
        
        asset_index = -1
        for i, (category, sub_category, component) in enumerate(self.app.ASSET_CATEGORIES):
            if component == asset_name:
                asset_index = i
                break
        
        if asset_index == -1:
            return "", ""
        
        asset_key = f"{asset_index + 1}_probability"
        asset_assessment = self.app.asset_data[latest_key]
        
        if asset_key not in asset_assessment:
            return "", ""
        
        asset_data = asset_assessment[asset_key]
        
        try:
            # Calculate likelihood and impact using app methods
            likelihood_cat = self.app.get_asset_likelihood_for_key(asset_key)
            impact_cat = self.app.get_asset_impact_for_key(asset_key)
            
            if likelihood_cat >= 0:
                likelihood_cat = self.app.value_to_category(likelihood_cat)
            else:
                likelihood_cat = ""
                
            if impact_cat >= 0:
                impact_cat = self.app.value_to_category(impact_cat)
            else:
                impact_cat = ""
            
            return likelihood_cat, impact_cat
            
        except (ValueError, KeyError):
            return "", ""

    def _parse_word_document_simple(self, doc):
        """Parse Word document for import (simplified version)"""
        try:
            tables = doc.tables
            
            if len(tables) == 0:
                messagebox.showwarning("Warning", "No tables found in document")
                return
            
            # Extract threat names and parse tables
            all_threat_names = self._extract_threat_names_from_document(doc)
            logging.info(f"Found threat names in document: {all_threat_names}")
            
            for threat_name in all_threat_names:
                if threat_name not in self.app.threat_data:
                    self.app.threat_data[threat_name] = {}
            
            # Track parsing statistics
            asset_tables_found = 0
            threat_tables_found = 0
            current_threat_name = None  # Track current threat context
            
            # Parse tables
            for i, table in enumerate(tables):
                if len(table.rows) == 0:
                    continue
                    
                header_texts = []
                if len(table.rows) > 0:
                    for cell in table.rows[0].cells:
                        header_texts.append(cell.text.strip().lower())
                
                # Log headers for debugging
                logging.info(f"Table {i} headers: {header_texts[:5]}...")  # First 5 headers
                
                # Skip controls tables (6 columns with Control Title, Control ID, etc.)
                if ('control title' in header_texts and 'control id' in header_texts and 
                    'description' in header_texts and len(header_texts) == 6):
                    logging.info(f"Skipping controls table at index {i}")
                    # Reset threat context after controls table - next threat table needs new threat name
                    current_threat_name = None
                    logging.info("Reset threat context after controls table")
                    continue
                
                # Asset assessment table (15 columns)
                if ('category' in header_texts and 'sub-category' in header_texts and 
                    'asset' in header_texts and 'likelihood' in header_texts):
                    logging.info(f"Found asset assessment table (table {i})")
                    self._parse_asset_table_simple(table)
                    asset_tables_found += 1
                    continue
                
                # Threat risk assessment table (22 columns) - more specific recognition
                has_asset_col = 'asset' in header_texts
                has_threat_likelihood = any('threat likelihood' in h for h in header_texts)
                has_threat_impact = any('threat impact' in h for h in header_texts)
                has_threat_risk = any('threat risk' in h for h in header_texts)
                has_a_prefix = any(h.startswith('a-') for h in header_texts)
                has_t_prefix = any(h.startswith('t-') for h in header_texts)
                
                # Additional check: make sure this is not a controls table
                is_controls_table = ('control title' in header_texts or 'control id' in header_texts)
                
                # Check for threat table pattern
                is_threat_table = (has_asset_col and has_threat_likelihood and has_threat_impact and 
                                   (has_threat_risk or has_a_prefix or has_t_prefix) and 
                                   not is_controls_table)
                
                if is_threat_table:
                    # Always look for threat name for each threat table, passing the known threat names
                    threat_name = self._find_threat_name_for_table(doc, i, all_threat_names)
                    if threat_name:
                        current_threat_name = threat_name  # Update current context
                        logging.info(f"Found threat risk assessment table for '{threat_name}' (table {i})")
                        logging.info(f"Table headers: {header_texts[:10]}...")  # More headers for debugging
                        self._parse_threat_table_simple(threat_name, table)
                        threat_tables_found += 1
                    else:
                        logging.warning(f"Threat table found but no threat name identified (table {i})")
                        logging.warning(f"Headers were: {header_texts}")
                        logging.warning(f"Current threat context was: {current_threat_name}")
                    continue
                
                # Log unrecognized tables for debugging
                logging.info(f"Unrecognized table {i} with {len(header_texts)} columns: {header_texts[:8]}...")
            
            # Log final statistics
            logging.info(f"Import parsing completed: {asset_tables_found} asset tables, {threat_tables_found} threat tables processed")
            
            # Log final imported data summary
            total_threats = len(self.app.threat_data)
            total_threat_asset_combinations = sum(len(threat_data) for threat_data in self.app.threat_data.values())
            
            logging.info(f"Final import summary:")
            logging.info(f"  - {total_threats} distinct threats imported")
            logging.info(f"  - {total_threat_asset_combinations} threat-asset combinations")
            
            for threat_name, threat_data in self.app.threat_data.items():
                logging.info(f"  - Threat '{threat_name}': {len(threat_data)} assets assessed")
                        
        except Exception as e:
            logging.error(f"Error parsing document: {str(e)}")
            raise

    def _extract_threat_names_from_document(self, doc):
        """Extract threat names from document paragraphs using known threats from Threat.csv"""
        threat_names = []
        
        # Load known threats from Threat.csv
        known_threats = set()
        threats_file = os.path.join(get_base_path(), "Threat.csv")
        try:
            with open(threats_file, 'r', newline='', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile, delimiter=';')
                for row in reader:
                    threat_name = row.get('THREAT', '').strip()
                    if threat_name:
                        known_threats.add(threat_name)
            logging.info(f"Loaded {len(known_threats)} known threats from Threat.csv")
        except Exception as e:
            logging.error(f"Error loading threats from Threat.csv: {e}")
            # Fallback to empty set if can't load file
            known_threats = set()
        
        # Extract threat names from document paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Method 1: Direct pattern "Risk Assessment for [threat_name]"
            if 'Risk Assessment for' in text:
                threat_name = text.replace('Risk Assessment for', '').strip()
                if threat_name and threat_name not in threat_names:
                    # Verify this is a known threat
                    if threat_name in known_threats:
                        threat_names.append(threat_name)
                        logging.info(f"Extracted threat name via 'Risk Assessment for': {threat_name}")
                    else:
                        logging.warning(f"Unknown threat name found: {threat_name}")
                continue
            
            # Method 2: Look for exact matches with known threats FIRST
            # This will catch headings like "Abuse of leaked data", "Denial of Service (DoS)"
            if text in known_threats and text not in threat_names:
                threat_names.append(text)
                logging.info(f"Extracted known threat name: {text}")
                continue
                
            # Method 3: Check for partial matches (case-insensitive) with known threats
            if text and len(text) > 3:  # Skip very short texts
                for known_threat in known_threats:
                    # Exact match (case-insensitive)
                    if text.lower() == known_threat.lower() and text not in threat_names:
                        threat_names.append(known_threat)  # Use the canonical form from CSV
                        logging.info(f"Extracted threat name via case-insensitive match: {known_threat}")
                        break
                    # Check if paragraph contains the threat name
                    elif known_threat.lower() in text.lower() and known_threat not in threat_names:
                        # Additional check: make sure it's not part of a larger sentence
                        if len(text.split()) <= 10:  # Likely a heading, not a sentence
                            threat_names.append(known_threat)
                            logging.info(f"Extracted threat name via partial match: {known_threat}")
                            break
        
        logging.info(f"Total threat names extracted: {len(threat_names)} - {threat_names}")
        return threat_names

    def _find_threat_name_for_table(self, doc, table_index, known_threat_names=None):
        """
        Find threat name for a table based on the specific structure of our generated reports.
        Structure in Detailed Threat Analysis:
        1. Threat Name (heading level 2)
        2. Risk Assessment for Threat Name (heading level 3) 
        3. Threat table (22 columns)
        4. Mitigation Controls for Threat Name (heading level 3)
        5. Controls table (6 columns)
        
        Args:
            doc: Word document
            table_index: Index of the table to find threat name for
            known_threat_names: List of threat names found at the beginning of the document
        """
        try:
            # Load known threats from Threat.csv for validation
            known_threats = set()
            threats_file = os.path.join(get_base_path(), "Threat.csv")
            try:
                with open(threats_file, 'r', newline='', encoding='utf-8') as csvfile:
                    reader = csv.DictReader(csvfile, delimiter=';')
                    for row in reader:
                        threat_name = row.get('THREAT', '').strip()
                        if threat_name:
                            known_threats.add(threat_name)
            except Exception as e:
                logging.error(f"Error loading threats for validation: {e}")
                known_threats = set()
            
            # Also add known threat names from document scan to our validation set
            if known_threat_names:
                for threat_name in known_threat_names:
                    known_threats.add(threat_name)
                logging.info(f"Added {len(known_threat_names)} threat names from document scan")
            
            # Find the target table and get the paragraphs before it
            table_count = 0
            current_table = None
            
            # Iterate through all tables in the document
            for table in doc.tables:
                if table_count == table_index:
                    current_table = table
                    break
                table_count += 1
            
            if not current_table:
                logging.warning(f"Table {table_index} not found")
                return None
            
            # Get the table element
            table_element = current_table._element
            
            # Look for paragraphs before this table
            parent = table_element.getparent()
            table_position = list(parent).index(table_element)
            
            # Collect paragraphs before the table (skip table elements)
            paragraphs_before = []
            for i in range(table_position - 1, -1, -1):
                element = parent[i]
                
                # Check if it's a paragraph
                if element.tag.endswith('p'):
                    # Extract text from paragraph using the most reliable method
                    paragraph_text = ""
                    
                    # Method 1: Try to get text from text elements directly using xpath
                    try:
                        for text_elem in element.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            if text_elem.text:
                                paragraph_text += text_elem.text
                    except:
                        pass
                    
                    # Method 2: If no text found, iterate through all text nodes
                    if not paragraph_text.strip():
                        def extract_text_recursive(elem):
                            text = ""
                            if elem.text:
                                text += elem.text
                            for child in elem:
                                text += extract_text_recursive(child)
                            if elem.tail:
                                text += elem.tail
                            return text
                        
                        paragraph_text = extract_text_recursive(element)
                    
                    # Clean up the text
                    paragraph_text = paragraph_text.strip()
                    # Remove any repeated spaces or newlines
                    import re
                    paragraph_text = re.sub(r'\s+', ' ', paragraph_text)
                    
                    if paragraph_text:  # Only add non-empty paragraphs
                        paragraphs_before.append(paragraph_text)
                        logging.info(f"Found paragraph {len(paragraphs_before)} before table {table_index}: '{paragraph_text}'")
                    
                    # Stop after finding 3 paragraphs (to be safe)
                    if len(paragraphs_before) >= 3:
                        break
                
                # If we encounter another table, we've gone too far
                elif element.tag.endswith('tbl'):
                    # Continue to skip other tables but don't count them
                    continue
            
            # Helper function to find best matching threat from known threats
            def find_best_threat_match(found_text, known_threat_names_list):
                """Find the best matching threat from the known threats list"""
                if not found_text or not known_threat_names_list:
                    return found_text
                
                found_text_lower = found_text.lower()
                
                # Method 1: Check if any known threat is a substring of what we found
                for known_threat in known_threat_names_list:
                    if known_threat.lower() in found_text_lower:
                        logging.info(f"Found substring match: '{found_text}' contains '{known_threat}'")
                        return known_threat
                
                # Method 2: Check if what we found is a substring of any known threat
                for known_threat in known_threat_names_list:
                    if found_text_lower in known_threat.lower():
                        logging.info(f"Found reverse substring match: '{known_threat}' contains '{found_text}'")
                        return known_threat
                
                # Method 3: Fuzzy matching for similar strings
                best_match = None
                best_similarity = 0
                for known_threat in known_threat_names_list:
                    # Simple similarity check based on common words
                    found_words = set(found_text_lower.split())
                    known_words = set(known_threat.lower().split())
                    
                    if found_words and known_words:
                        common_words = found_words.intersection(known_words)
                        similarity = len(common_words) / max(len(found_words), len(known_words))
                        
                        if similarity > best_similarity and similarity >= 0.5:  # At least 50% similarity
                            best_similarity = similarity
                            best_match = known_threat
                
                if best_match:
                    logging.info(f"Found fuzzy match: '{found_text}' -> '{best_match}' (similarity: {best_similarity:.2f})")
                    return best_match
                
                return found_text
            
            # Now analyze the paragraphs to find the threat name
            # Structure should be:
            # paragraphs_before[0] = "Risk Assessment for Threat Name" (immediately before table)
            # paragraphs_before[1] = "Threat Name" (the actual threat name heading)
            
            if len(paragraphs_before) >= 1:
                # Check the first paragraph (immediately before table)
                first_paragraph = paragraphs_before[0]
                logging.info(f"Checking paragraph immediately before table {table_index}: '{first_paragraph}'")
                
                # Method 1: Look for "Risk Assessment for [threat_name]" pattern (this should be our primary match)
                if first_paragraph.startswith('Risk Assessment for'):
                    threat_name = first_paragraph.replace('Risk Assessment for', '').strip()
                    logging.info(f"Extracted threat name from 'Risk Assessment for' pattern: '{threat_name}'")
                    
                    # Try to find the best match from known threats
                    if known_threat_names:
                        matched_threat = find_best_threat_match(threat_name, known_threat_names)
                        if matched_threat != threat_name:
                            logging.info(f"Corrected threat name from '{threat_name}' to '{matched_threat}'")
                            return matched_threat
                    
                    # Validate against known threats
                    if threat_name in known_threats:
                        logging.info(f"Found valid threat via 'Risk Assessment for' pattern: {threat_name}")
                        return threat_name
                    else:
                        # Try case-insensitive match
                        for known_threat in known_threats:
                            if threat_name.lower() == known_threat.lower():
                                logging.info(f"Found valid threat via case-insensitive 'Risk Assessment for' pattern: {known_threat}")
                                return known_threat
                        
                        # If not found in known threats, still return it (CSV might be incomplete)
                        logging.info(f"Using extracted threat name (not in CSV): {threat_name}")
                        return threat_name
            
            # If we have more paragraphs, check the second one (should be the direct threat name)
            if len(paragraphs_before) >= 2:
                second_paragraph = paragraphs_before[1]
                logging.info(f"Checking 2nd paragraph before table {table_index}: '{second_paragraph}'")
                
                # Try to find the best match from known threats first
                if known_threat_names:
                    matched_threat = find_best_threat_match(second_paragraph, known_threat_names)
                    if matched_threat != second_paragraph:
                        logging.info(f"Corrected threat name from '{second_paragraph}' to '{matched_threat}'")
                        return matched_threat
                
                # Method 2: Direct threat name match
                if second_paragraph in known_threats:
                    logging.info(f"Found exact threat match in 2nd paragraph: {second_paragraph}")
                    return second_paragraph
                
                # Method 3: Case-insensitive exact match
                for known_threat in known_threats:
                    if second_paragraph.lower() == known_threat.lower():
                        logging.info(f"Found case-insensitive exact match in 2nd paragraph: {known_threat}")
                        return known_threat
                
                # Method 4: Check if paragraph contains a known threat name (for short paragraphs only)
                if len(second_paragraph.split()) <= 8:  # Only for reasonably short headings
                    for known_threat in known_threats:
                        if known_threat.lower() in second_paragraph.lower():
                            logging.info(f"Found threat via substring match in 2nd paragraph: {known_threat}")
                            return known_threat
                
                # Method 5: If we found a heading-like text but it's not in known threats,
                # it might still be a valid threat name (in case CSV is incomplete)
                if (len(second_paragraph.split()) <= 6 and 
                    not any(word in second_paragraph.lower() for word in ['table', 'assessment', 'control', 'overview', 'mitigation'])):
                    logging.info(f"Found potential threat name (not in CSV) in 2nd paragraph: {second_paragraph}")
                    return second_paragraph
            
            logging.warning(f"No threat name found for table {table_index}. Found {len(paragraphs_before)} paragraphs before table.")
            if paragraphs_before:
                logging.warning(f"Paragraphs were: {paragraphs_before}")
            return None
            
        except Exception as e:
            logging.error(f"Error finding threat name for table {table_index}: {str(e)}")
            return None

    def _parse_asset_table_simple(self, table):
        """Parse asset assessment table for import - matches export format exactly"""
        try:
            if len(table.rows) < 2:
                return
            
            # Expected format: 15 columns
            # Category(0), Sub-category(1), Asset(2), 9 criteria(3-11), Likelihood(12), Impact(13), Risk(14)
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            asset_key = f"imported_{timestamp}"
            self.app.asset_data[asset_key] = {}
            
            for row_idx in range(1, len(table.rows)):
                row = table.rows[row_idx]
                cells = [cell.text.strip() for cell in row.cells]
                
                if len(cells) < 15:  # Must have all 15 columns
                    continue
                
                asset_name = cells[2]  # Asset name is in column 2
                
                if not asset_name or asset_name in ["", "N/A"]:
                    continue
                
                # Find asset index by name in current system
                asset_index = -1
                for i, (category, sub_category, component) in enumerate(self.app.ASSET_CATEGORIES):
                    if component.lower() == asset_name.lower():
                        asset_index = i
                        break
                
                if asset_index < 0:
                    continue
                
                probability_key = f"{asset_index + 1}_probability"
                
                # Extract criteria values from columns 3-11 (9 criteria)
                asset_data = {}
                for i in range(9):
                    criteria_col = 3 + i
                    criteria_value = cells[criteria_col].strip()
                    if criteria_value and criteria_value.isdigit() and 1 <= int(criteria_value) <= 5:
                        asset_data[str(i)] = criteria_value
                
                # Only save if we have at least some criteria data
                if asset_data:
                    self.app.asset_data[asset_key][probability_key] = asset_data
                    logging.info(f"Imported asset data for {asset_name}: {len(asset_data)} criteria")
                                
        except Exception as e:
            logging.error(f"Error parsing asset table: {str(e)}")
            import traceback
            traceback.print_exc()

    def _parse_threat_table_simple(self, threat_name, table):
        """Parse threat table for import - matches export format exactly"""
        try:
            if len(table.rows) < 2:
                logging.warning(f"Threat table for {threat_name} has insufficient rows")
                return
            
            if threat_name not in self.app.threat_data:
                self.app.threat_data[threat_name] = {}
            
            # Check number of columns
            num_cols = len(table.rows[0].cells) if len(table.rows) > 0 else 0
            logging.info(f"Processing threat table for '{threat_name}' with {num_cols} columns")
            
            # Expected format: 22 columns
            # Asset(0), 9 Asset criteria(1-9), Asset L/I(10-11), 7 Threat criteria(12-18), Threat L/I/R(19-21)
            
            assets_processed = 0
            for row_idx in range(1, len(table.rows)):
                row = table.rows[row_idx]
                cells = [cell.text.strip() for cell in row.cells]
                
                if len(cells) < 12:  # Need at least up to threat criteria start
                    logging.warning(f"Row {row_idx} has only {len(cells)} cells, skipping")
                    continue
                
                asset_name = cells[0]  # Asset name is in column 0
                
                if asset_name in ["", "N/A", "No risk assessment data available"]:
                    continue
                
                if not asset_name:
                    continue
                
                # Find asset index by name in current system
                asset_index = -1
                for i, (category, sub_category, component) in enumerate(self.app.ASSET_CATEGORIES):
                    if component.lower() == asset_name.lower():
                        asset_index = i
                        break
                
                if asset_index < 0:
                    logging.warning(f"Asset '{asset_name}' not found in categories")
                    continue
                
                asset_key = f"{asset_index + 1}_probability"
                
                # Extract threat criteria from available columns
                threat_data = {}
                criteria_start = 12 if num_cols >= 22 else max(0, len(cells) - 7)  # Adaptive start
                
                for i in range(7):
                    criteria_col = criteria_start + i
                    if criteria_col < len(cells):
                        criteria_value = cells[criteria_col].strip()
                        if criteria_value and criteria_value.replace('.', '').isdigit():
                            try:
                                val = int(float(criteria_value))
                                if 1 <= val <= 5:
                                    threat_data[str(i)] = str(val)
                            except (ValueError, TypeError):
                                pass
                
                # Save if we have at least some threat criteria data
                if threat_data:
                    self.app.threat_data[threat_name][asset_key] = threat_data
                    assets_processed += 1
                    logging.info(f"Imported threat data for {threat_name}-{asset_name}: {len(threat_data)} criteria")
            
            logging.info(f"Processed {assets_processed} assets for threat '{threat_name}'")
                                
        except Exception as e:
            logging.error(f"Error parsing threat table for {threat_name}: {str(e)}")
            import traceback
            traceback.print_exc()

    def _create_synthetic_threat_data(self, likelihood_cat, impact_cat):
        """Create synthetic threat data from categories"""
        try:
            category_to_value = {
                "Very Low": 1, "Low": 2, "Medium": 3, "High": 4, "Very High": 5
            }
            
            likelihood_val = category_to_value.get(likelihood_cat, 0)
            impact_val = category_to_value.get(impact_cat, 0)
            
            if likelihood_val == 0 or impact_val == 0:
                return None
            
            synthetic_data = {}
            
            # 5 criteria for likelihood
            for i in range(5):
                synthetic_data[str(i)] = str(likelihood_val)
            
            # 2 criteria for impact
            for i in range(5, 7):
                synthetic_data[str(i)] = str(impact_val)
            
            return synthetic_data
            
        except Exception as e:
            logging.error(f"Error creating synthetic threat data: {str(e)}")
            return None

    def _create_synthetic_asset_data(self, likelihood_cat, impact_cat):
        """Create synthetic asset data from categories"""
        try:
            category_to_value = {
                "Very Low": 1, "Low": 2, "Medium": 3, "High": 4, "Very High": 5
            }
            
            likelihood_val = category_to_value.get(likelihood_cat, 0)
            impact_val = category_to_value.get(impact_cat, 0)
            
            if likelihood_val == 0 or impact_val == 0:
                return None
            
            synthetic_data = {}
            
            # 4 criteria for likelihood
            for i in range(4):
                synthetic_data[str(i)] = str(likelihood_val)
            
            # 5 criteria for impact
            for i in range(4, 9):
                synthetic_data[str(i)] = str(impact_val)
            
            return synthetic_data
            
        except Exception as e:
            logging.error(f"Error creating synthetic asset data: {str(e)}")
            return None

    def test_import_with_existing_file(self):
        """Test import functionality with an existing Word report"""
        try:
            # Find the most recent report file
            import glob
            pattern = os.path.join(get_base_path(), "Risk_Assessment_Report_*.docx")
            report_files = glob.glob(pattern)
            
            if not report_files:
                logging.warning("No existing report files found for testing")
                return False
            
            # Use the most recent file
            latest_file = max(report_files, key=os.path.getctime)
            
            # Save current state
            original_asset_data = self.app.asset_data.copy()
            original_threat_data = self.app.threat_data.copy()
            
            try:
                # Clear data to test clean import
                self.app.asset_data = {}
                self.app.threat_data = {}
                
                # Import the file
                doc = Document(latest_file)
                self._parse_word_document_simple(doc)
                
                # Check results
                imported_threats = len(self.app.threat_data)
                imported_assets = 0
                threat_asset_combinations = 0
                
                if self.app.asset_data:
                    for key, data in self.app.asset_data.items():
                        if key.startswith('imported_'):
                            imported_assets = len([k for k in data.keys() if k.endswith('_probability')])
                            break
                
                for threat_name, threat_data in self.app.threat_data.items():
                    threat_asset_combinations += len(threat_data)
                
                success = imported_threats > 0 or imported_assets > 0
                
                if success:
                    logging.info(f"Import test successful: {imported_threats} threats, {threat_asset_combinations} threat-asset combinations, {imported_assets} assets from {os.path.basename(latest_file)}")
                else:
                    logging.warning("Import test failed: no data imported")
                
                return success
                
            finally:
                # Always restore original state
                self.app.asset_data = original_asset_data
                self.app.threat_data = original_threat_data
                
        except Exception as e:
            logging.error(f"Error in import test: {str(e)}")
            return False