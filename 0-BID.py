# Risk Assessment Tool - BID phase
# Purpose: Calculate risk value of a ITT from category of the project
# Author: Thesis work for space program risk assessment tool Giuseppe Nonni 1948023 giuseppe.nonni@gmail.com

import tkinter as tk
from tkinter import messagebox, ttk
import tkinter.font as tkFont
from datetime import datetime
import os
import sys

def get_base_path():
    """Get the base path for the application (works with both .py and .exe)"""
    if getattr(sys, 'frozen', False):
        # Running as compiled executable
        return os.path.dirname(sys.executable)
    else:
        # Running as script
        return os.path.dirname(os.path.abspath(__file__))

try:
    from PIL import Image, ImageTk, ImageDraw
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class BIDOptimized:        
        # Save Risk Assessment Data
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    risk_filename = f"BID_Risk_Assessment_{timestamp}.csv"
        
        # Create Output directory if it doesn't exist
    output_dir = os.path.join(get_base_path(), "Output")
    os.makedirs(output_dir, exist_ok=True)
        
    risk_filepath = os.path.join(output_dir, risk_filename)   # Configuration data
    
    COLORS = {
        'primary': '#4a90c2', 'secondary': '#dc3545', 'success': '#28a745',
        'white': '#ffffff', 'light': '#f8f9fa', 'dark': '#2c3e50',
        'gray': '#6c757d', 'blue': '#007bff', 'green': '#d4edda',
        'yellow': '#fff3cd', 'red': '#f8d7da', 'dark_red': '#dc3545'
    }
    
    # Table 1 data (Score Matrix) - Removed Weight column
    TABLE1_DATA = [
        ["Category", "Score 1 (Low)", "Score 2 (Significative)", "Score 3 (Moderate)", "Score 4 (High)"],
        ["Cybersecurity Requirements", "Clear, specific CIA objectives and mapped controls", "Partial objectives, general security references", "Vague mention of cybersecurity without clear objectives", "No security objectives mentioned"],
        ["Security Architecture Constraints", "Defined secure architecture, protocols and constraints", "General reference to secure design without details", "Weak constraints, non-binding suggestions", "No architectural constraints present"],
        ["Cryptographic Requirements", "Detailed crypto specs (e.g. AES256, PKI), lifecycle defined", "Crypto required but not specified", "Crypto mentioned vaguely, unclear implementation strategy", "No mention of encryption or key management"],
        ["Authentication & Access Control", "Clear roles, access policies, identity/authentication methods", "Generic role-based access noted", "Some access control logic implied but no detail", "No mention of access control or identity management"],
        ["Supply Chain Security", "Trusted suppliers required, integrity checks mandated", "Supplier lists validated but not verified", "Open supplier selection, no trust/integrity verification", "No supply chain considerations present"],
        ["Threat Modeling Guidelines", "Threat model provided or referenced", "Reference to general threat types", "High-level mention of risk environment", "No threat modeling or attack surface identified"],
        ["Security Compliance References", "Full list of mandatory compliance standards", "Some standards listed, not mandatory", "Mentioned standards optional or vague", "No standards or frameworks referenced"],
        ["Security Validation Requirements", "Detailed validation strategy including scope and responsibility", "Validation required but not detailed", "Unclear expectations for testing/audits", "No mention of security validation"],
        ["Incident Response Expectations", "Incident response roles, deadlines, escalation paths defined", "Some response actions outlined", "Minimal requirements for incident handling", "No incident response planning mentioned"],
        ["Data Protection and Privacy", "Full compliance expectations and procedures included", "Compliance mentioned but procedures vague", "Compliance cited but not related to mission data", "No mention of data protection or privacy"],
        ["Cybersecurity Historical Data", "Documented past incidents and mitigation strategies provided", "General lessons learned included", "Incomplete data or single example used", "No historical data on cybersecurity issues"]
    ]
    
    # Table 2 data (Risk Assessment) - Keep Weight column
    TABLE2_DATA = [
        ["Category", "Value (1-4)", "Weight", "Inapplicability"],
        *[[row[0], "", "0.15" if row[0] == "Cybersecurity Requirements" else 
           "0.12" if row[0] == "Security Architecture Constraints" else
           "0.10" if row[0] == "Cryptographic Requirements" else
           "0.08" if row[0] == "Authentication & Access Control" else
           "0.12" if row[0] == "Supply Chain Security" else
           "0.08" if row[0] == "Threat Modeling Guidelines" else
           "0.07" if row[0] == "Security Compliance References" else
           "0.10" if row[0] == "Security Validation Requirements" else
           "0.05" if row[0] == "Incident Response Expectations" else
           "0.07" if row[0] == "Data Protection and Privacy" else
           "0.06" if row[0] == "Cybersecurity Historical Data" else "0.00", ""] for row in TABLE1_DATA[1:]]
    ]
    
    # Table 3 data (Results)
    TABLE3_DATA = [
        ["Total Score", "Risk Level"], 
        ["0.000", "Very Low"],
        ["Score Range", "Level"],  # Sub-header
        ["0-0.1", "Very Low"], 
        ["0.1-0.4", "Low"], 
        ["0.4-0.7", "Medium"], 
        ["0.7-0.9", "High"], 
        ["0.9-1", "Very High"]
    ]
    
    # Risk level colors
    RISK_COLORS = {
        'Very Low': ('#f8f9fa', '#6c757d'),
        'Low': ('#d4edda', '#155724'),
        'Medium': ('#fff3cd', '#856404'), 
        'High': ('#f8d7da', '#721c24'),
        'Very High': ('#dc3545', '#ffffff')
    }
    
    # Criteria descriptions for help window
    CRITERIA_DESCRIPTIONS = {
        "Cybersecurity Requirements": "Defines the confidentiality, integrity, and availability (CIA) objectives and security controls that must be implemented in the system.",
        "Security Architecture Constraints": "Specifies the secure design principles, protocols, and architectural constraints that guide the system's security implementation.",
        "Cryptographic Requirements": "Details the encryption standards, key management procedures, and cryptographic lifecycle requirements for data protection.",
        "Authentication & Access Control": "Establishes user identity verification methods, role-based access policies, and authorization mechanisms for system resources.",
        "Supply Chain Security": "Ensures the integrity and trustworthiness of suppliers, components, and third-party services throughout the supply chain.",
        "Threat Modeling Guidelines": "Provides systematic approach to identify, analyze, and mitigate potential security threats and attack vectors.",
        "Security Compliance References": "Lists mandatory security standards, frameworks, and regulatory requirements that must be adhered to.",
        "Security Validation Requirements": "Defines testing strategies, audit procedures, and validation methods to verify security implementation effectiveness.",
        "Incident Response Expectations": "Outlines procedures, roles, and timelines for detecting, responding to, and recovering from security incidents.",
        "Data Protection and Privacy": "Specifies requirements for protecting sensitive data and ensuring compliance with privacy regulations and policies.",
        "Cybersecurity Historical Data": "Documents past security incidents, lessons learned, and proven mitigation strategies relevant to the current context."
    }

    def __init__(self, root):
        self.root = root
        self.root.title("BID Phase")
        self.setup_scaling()
        self.setup_ui()
        self.create_tables()
        self.update_total_score()

    def setup_scaling(self):
        """Calculate scale factors based on screen resolution"""
        screen_area = self.root.winfo_screenwidth() * self.root.winfo_screenheight()
        self.scale_factor = max(0.6, min(2.5, (screen_area / (1920 * 1080)) ** 0.5))
        
        # Scaled dimensions
        self.scaled_font_size = max(9, int(11 * self.scale_factor))
        self.scaled_title_font = max(14, int(16 * self.scale_factor))
        self.scaled_button_font = max(10, int(12 * self.scale_factor))
        self.scaled_padding = max(5, int(8 * self.scale_factor))
        self.scaled_button_padding = max(15, int(20 * self.scale_factor))
        self.scaled_cell_height = max(20, int(25 * self.scale_factor))

    def disable_mousewheel_on_combobox(self, combo):
        """Disable mouse wheel on combobox to prevent accidental value changes while allowing page scroll"""
        def on_mousewheel(event):
            # Check if the combobox dropdown is open
            try:
                if combo.tk.call('ttk::combobox::PopdownIsVisible', combo):
                    # If dropdown is open, allow normal combobox behavior
                    return
                else:
                    # If dropdown is closed, scroll the page instead of changing value
                    self.canvas.yview_scroll(int(-1*(event.delta/120)), "units")
                    return "break"  # Prevent combobox value change
            except:
                # Fallback: scroll the page
                self.canvas.yview_scroll(int(-1*(event.delta/120)), "units")
                return "break"
        
        combo.bind("<MouseWheel>", on_mousewheel)

    def setup_ui(self):
        """Setup main UI structure"""
        # Set window size to accommodate 1700px content with padding
        self.root.geometry("1800x950")
        self.root.state('zoomed')
        self.root.configure(bg=self.COLORS['white'])
        
        # Header
        header = tk.Frame(self.root, bg=self.COLORS['light'], height=80)
        header.pack(fill='x')
        header.pack_propagate(False)
        
        tk.Label(header, text="BID Phase - Risk Assessment", 
                font=('Segoe UI', self.scaled_title_font + 2, 'bold'),
                bg=self.COLORS['light'], fg=self.COLORS['dark']).pack(pady=(20, 10))
        
        # Scrollable canvas
        self.canvas = tk.Canvas(self.root, bg=self.COLORS['white'], highlightthickness=0)
        scrollbar_v = tk.Scrollbar(self.root, orient="vertical", command=self.canvas.yview)
        scrollbar_h = tk.Scrollbar(self.root, orient="horizontal", command=self.canvas.xview)
        
        self.main_frame = tk.Frame(self.canvas, bg=self.COLORS['white'])
        self.main_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        
        self.canvas.create_window((0, 0), window=self.main_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=scrollbar_v.set, xscrollcommand=scrollbar_h.set)
        
        # Pack layout
        self.canvas.pack(side="left", fill="both", expand=True, padx=15, pady=10)
        scrollbar_v.pack(side="right", fill="y", padx=(0, 5))
        scrollbar_h.pack(side="bottom", fill="x", padx=(15, 17), pady=(0, 5))
        
        # Mouse wheel scrolling
        self.canvas.bind_all("<MouseWheel>", lambda e: self.canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
        self.canvas.bind_all("<Shift-MouseWheel>", lambda e: self.canvas.xview_scroll(int(-1*(e.delta/120)), "units"))

    def create_modern_entry(self, parent, **kwargs):
        """Create modern styled entry widget"""
        defaults = {
            'relief': 'flat', 'bd': 0, 'highlightthickness': 2,
            'highlightcolor': self.COLORS['primary'],
            'highlightbackground': '#e1e5e9',
            'insertbackground': '#495057'
        }
        defaults.update(kwargs)
        
        entry = tk.Entry(parent, **defaults)
        entry.bind('<FocusIn>', lambda e: entry.configure(highlightcolor='#357abd'))
        entry.bind('<FocusOut>', lambda e: entry.configure(highlightcolor=self.COLORS['primary']))
        return entry

    def create_table(self, parent, title, data, color, interactive=False):
        """Generic table creation function"""
        container = tk.Frame(parent, bg=self.COLORS['white'])
        
        frame = tk.LabelFrame(container, text=title,
                             font=('Segoe UI', self.scaled_font_size + 2, 'bold'),
                             bg=self.COLORS['white'], fg=color, padx=25, pady=20,
                             relief='ridge', bd=2, labelanchor='n')
        frame.pack(fill='both', expand=True)
        
        cells = []
        combo_vars, check_vars = [], []
        
        for i, row in enumerate(data):
            cell_row, combo_row, check_row = [], [], []
            
            for j, text in enumerate(row):
                # Handle interactive elements for table 2
                if interactive and i > 0 and j == 1:  # ComboBox column
                    combo_var = tk.StringVar(value='1')
                    combo = ttk.Combobox(frame, textvariable=combo_var, values=['1', '2', '3', '4'],
                                       state='readonly', width=8, font=('Segoe UI', self.scaled_font_size - 1))
                    combo.bind('<<ComboboxSelected>>', self.update_total_score)
                    self.disable_mousewheel_on_combobox(combo)  # Prevent accidental value changes
                    combo.grid(row=i, column=j, padx=3, pady=3, sticky='ew', ipady=7)
                    cell_row.append(combo)
                    combo_row.append(combo_var)
                    check_row.append(None)
                    
                elif interactive and i > 0 and j == 2:  # Weight column
                    weight_cell = self.create_table_cell(frame, i, j, text, color, title)
                    weight_cell.grid(row=i, column=j, padx=3, pady=self.get_cell_pady(i, title), 
                                   sticky='ew', ipady=self.get_cell_ipady(i, title))
                    cell_row.append(weight_cell)
                    combo_row.append(None)
                    check_row.append(None)
                    
                elif interactive and i > 0 and j == 3:  # Checkbox column (back to j == 3)
                    check_var = tk.BooleanVar()
                    check_frame = tk.Frame(frame, bg=self.COLORS['white'], height=self.scaled_cell_height + 16)
                    check_frame.pack_propagate(False)
                    check_frame.grid(row=i, column=j, padx=3, pady=3, sticky='ew')
                    
                    checkbox = tk.Checkbutton(check_frame, variable=check_var, bg=self.COLORS['white'],
                                            activebackground=self.COLORS['white'], selectcolor=self.COLORS['white'],
                                            fg=self.COLORS['dark'], highlightthickness=0, relief='flat',
                                            command=lambda idx=i: self.update_weight_value(idx))
                    checkbox.place(relx=0.5, rely=0.5, anchor='center')
                    cell_row.append(check_frame)
                    combo_row.append(None)
                    check_row.append(check_var)
                    
                else:  # Regular entry cells
                    cell = self.create_table_cell(frame, i, j, text, color, title)
                    cell.grid(row=i, column=j, padx=3, pady=self.get_cell_pady(i, title), 
                             sticky='ew', ipady=self.get_cell_ipady(i, title))
                    cell_row.append(cell)
                    combo_row.append(None)
                    check_row.append(None)
            
            cells.append(cell_row)
            combo_vars.append(combo_row)
            check_vars.append(check_row)
        
        # Configure grid weights
        for j in range(len(data[0])):
            frame.grid_columnconfigure(j, weight=1)
        for i in range(len(data)):
            frame.grid_rowconfigure(i, minsize=self.scaled_cell_height)
            
        # Auto-resize columns
        self.auto_resize_table(cells, frame, len(data[0]), title == "Risk Assessment")
        
        return container, cells, combo_vars, check_vars
    
    def create_table_cell(self, parent, i, j, text, color, table_title):
        """Create individual table cell with appropriate styling"""
        # Determine cell styling based on position and table type
        is_header = i == 0 or (table_title == "Results" and i == 2)
        is_calculation_row = table_title == "Results" and i == 1
        
        # Check if this cell needs text wrapping (Score Matrix description columns)
        needs_wrapping = (table_title == "Score Matrix" and j in [1, 2, 3, 4] and len(text) > 50)
        
        if needs_wrapping:
            # Use Text widget for multiline text in Score Matrix description columns
            if is_header:
                text_widget = tk.Text(parent, bg=color, fg=self.COLORS['white'],
                                    font=('Segoe UI', self.scaled_font_size, 'bold'),
                                    wrap='word', relief='flat', bd=0, padx=5, pady=3,
                                    height=2, width=30, state='normal',
                                    highlightthickness=2, highlightcolor=self.COLORS['primary'],
                                    highlightbackground='#e1e5e9', insertbackground=self.COLORS['white'])
                text_widget.insert('1.0', text)
                text_widget.config(state='disabled')
                text_widget.tag_configure("center", justify='center')
                text_widget.tag_add("center", "1.0", "end")
            else:
                bg_color, fg_color = self.get_data_cell_colors(i, j, text, table_title)
                text_widget = tk.Text(parent, bg=bg_color, fg=fg_color,
                                    font=('Segoe UI', self.scaled_font_size - 1, 'normal'),
                                    wrap='word', relief='flat', bd=0, padx=5, pady=3,
                                    height=2, width=30, state='normal',
                                    highlightthickness=2, highlightcolor=self.COLORS['primary'],
                                    highlightbackground='#e1e5e9', insertbackground=fg_color)
                text_widget.insert('1.0', text)
                text_widget.config(state='disabled')
                # Configure left alignment for Score Matrix data cells
                text_widget.tag_configure("left", justify='left')
                text_widget.tag_add("left", "1.0", "end")
            return text_widget
        
        elif is_header:
            # Headers: All tables use center alignment
            justify = 'center'
            cell = self.create_modern_entry(parent, readonlybackground=color, state='normal',
                                          font=('Consolas' if i == 0 else 'Segoe UI', self.scaled_font_size, 'bold'),
                                          justify=justify, fg=self.COLORS['white'], insertbackground=self.COLORS['white'],
                                          highlightthickness=0)
            cell.insert(0, text)
            cell.configure(state='readonly')
            return cell
        
        elif is_calculation_row:
            cell = self.create_modern_entry(parent, readonlybackground=self.COLORS['white'], state='normal',
                                          font=('Segoe UI', self.scaled_font_size, 'bold'), justify='center',
                                          fg=self.COLORS['dark'], insertbackground=self.COLORS['dark'],
                                          highlightthickness=2, highlightcolor=color)
            cell.insert(0, text)
            cell.configure(state='readonly')
            return cell
        
        else:
            # Data cell styling
            bg_color, fg_color = self.get_data_cell_colors(i, j, text, table_title)
            font_style = ('Segoe UI', self.scaled_font_size - 1, 'bold' if j == 0 else 'normal')
              # Alignment rules for different tables
            if table_title == "Score Matrix":
                # Score Matrix: all columns left-aligned
                justify = 'left'
            elif table_title == "Results":
                # Results table: all columns center-aligned
                justify = 'center'
            else:
                # Risk Assessment: first column left, others center
                justify = 'left' if j == 0 else 'center'
                
            cell = self.create_modern_entry(parent, readonlybackground=bg_color, state='normal',
                                          font=font_style, justify=justify, fg=fg_color,
                                          insertbackground=fg_color, highlightthickness=1, highlightcolor=color)
            
            # Insert text and make readonly
            cell.insert(0, text)
            cell.configure(state='readonly')
            return cell

    def get_data_cell_colors(self, i, j, text, table_title):
        """Get appropriate colors for data cells"""
        if table_title == "Results" and j == 1 and i >= 3:  # Risk level colors
            level = text
            return self.RISK_COLORS.get(level, (self.COLORS['light'], self.COLORS['dark']))
        elif j == 0:  # First column
            return (self.COLORS['light'] if i % 2 == 0 else self.COLORS['white'], self.COLORS['dark'])
        else:
            return (self.COLORS['white'], '#495057')

    def get_cell_pady(self, i, table_title):
        """Get appropriate pady for cell"""
        if table_title == "Results":
            return 1 if i <= 1 else 2
        return 3

    def get_cell_ipady(self, i, table_title):
        """Get appropriate ipady for cell"""
        if table_title == "Results":
            return 6 if i <= 1 else (8 if i == 2 else 7)
        return 10

    def create_tables(self):
        """Create all three tables"""
        # Table 1 - Score Matrix
        table1_container, self.cells1, _, _ = self.create_table(
            self.main_frame, "Score Matrix", self.TABLE1_DATA, self.COLORS['dark'])
        table1_container.pack(fill='both', expand=True, padx=20, pady=(20, 15))
        
        # Bottom frame for tables 2 and 3
        bottom_frame = tk.Frame(self.main_frame, bg=self.COLORS['white'])
        bottom_frame.pack(fill='both', expand=True, padx=20, pady=15)
        bottom_frame.grid_columnconfigure(0, weight=2)
        bottom_frame.grid_columnconfigure(1, weight=1)
        bottom_frame.grid_rowconfigure(0, weight=1)
        
        # Table 2 - Risk Assessment (interactive)
        table2_container, self.cells2, self.combo_vars, self.check_vars = self.create_table(
            bottom_frame, "Risk Assessment", self.TABLE2_DATA, self.COLORS['secondary'], interactive=True)
        table2_container.grid(row=0, column=0, sticky='nsew', padx=(0, 15), pady=15)
        
        # Table 3 - Results
        table3_container, self.cells3, _, _ = self.create_table(
            bottom_frame, "Results", self.TABLE3_DATA, self.COLORS['success'])
        table3_container.grid(row=0, column=1, sticky='nsew', padx=(10, 0), pady=15)
        
        # Save button
        self.create_save_button()

    def create_save_button(self):
        """Create save button and help button with hover effects"""
        button_frame = tk.Frame(self.main_frame, bg=self.COLORS['white'])
        button_frame.pack(pady=30, padx=20)
        
        button_container = tk.Frame(button_frame, bg=self.COLORS['light'], relief='ridge', bd=1)
        button_container.pack()
        
        # Container for buttons side by side
        buttons_inner_frame = tk.Frame(button_container, bg=self.COLORS['light'])
        buttons_inner_frame.pack(padx=3, pady=3)
        
        # Save button
        self.save_button = tk.Button(buttons_inner_frame, text="💾 Save Data", command=self.save_data,
                                   font=('Segoe UI', self.scaled_button_font + 1, 'bold'),
                                   bg=self.COLORS['blue'], fg=self.COLORS['white'],
                                   activebackground='#0056b3', activeforeground=self.COLORS['white'],
                                   relief='flat', bd=0, cursor='hand2',
                                   padx=self.scaled_button_padding + 15, pady=self.scaled_padding + 8)
        self.save_button.pack(side='left', padx=(0, 10))
        
        # Help button
        self.help_button = tk.Button(buttons_inner_frame, text="❓ Help", command=self.show_help,
                                   font=('Segoe UI', self.scaled_button_font + 1, 'bold'),
                                   bg=self.COLORS['gray'], fg=self.COLORS['white'],
                                   activebackground='#545b62', activeforeground=self.COLORS['white'],
                                   relief='flat', bd=0, cursor='hand2',
                                   padx=self.scaled_button_padding + 10, pady=self.scaled_padding + 8)
        self.help_button.pack(side='left')
        
        # Hover effects
        self.save_button.bind("<Enter>", lambda e: self.save_button.config(bg='#0056b3'))
        self.save_button.bind("<Leave>", lambda e: self.save_button.config(bg=self.COLORS['blue']))
        
        self.help_button.bind("<Enter>", lambda e: self.help_button.config(bg='#545b62'))
        self.help_button.bind("<Leave>", lambda e: self.help_button.config(bg=self.COLORS['gray']))

    def auto_resize_table(self, cells, frame, num_cols, is_risk_table=False):
        """Auto-resize table columns based on content using original copy widths"""
        import tkinter.font as tkFont
        col_widths = []
        
        for j in range(num_cols):
            max_width = 0
            
            # Calculate maximum width for this column
            for i in range(len(cells)):
                if isinstance(cells[i][j], tk.Entry):
                    cell = cells[i][j]
                    text = cell.get() if hasattr(cell, 'get') else ""
                    font = cell.cget("font") if hasattr(cell, 'cget') else ("Segoe UI", self.scaled_font_size)
                else:
                    # Skip non-Entry widgets (ComboBox, Checkboxes, etc.)
                    continue
                
                if text:
                    try:
                        font_obj = tkFont.Font(font=font)
                        text_width = font_obj.measure(text)
                        max_width = max(max_width, text_width)
                    except:
                        # Fallback for font issues
                        max_width = max(max_width, len(text) * 8)
            
            # Special handling for different tables and columns
            if frame.master.master and hasattr(frame.master.master, 'winfo_children'):
                # Get table title from parent structure
                table_containers = frame.master.master.winfo_children()
                table_index = 0
                for widget in table_containers:
                    if hasattr(widget, 'winfo_children') and frame.master in widget.winfo_children():
                        break
                    table_index += 1
                
                # Apply specific width rules based on table and column (original copy widths)
                if table_index == 0:  # Score Matrix (without Weight column)
                    if j == 0:  # Category column
                        min_width = 200
                    elif j in [1, 2, 3, 4]:  # Score columns (1-4)
                        min_width = 350  # Same width as original for long descriptions
                elif is_risk_table:  # Risk Assessment table (with Weight column)
                    if j == 0:  # Category column
                        min_width = 250
                    elif j == 1:  # Value column
                        min_width = 100
                    elif j == 2:  # Weight column
                        min_width = 80
                    else:  # Inapplicability column
                        min_width = 120
                else:  # Results table
                    min_width = 150
            else:
                min_width = 150  # Default minimum width
            
            # Use the larger of calculated width or minimum width
            final_width = max(max_width + int(30 * self.scale_factor), 
                            int(min_width * self.scale_factor))
            col_widths.append(final_width)
        
        # Set column configurations
        for j in range(num_cols):
            frame.grid_columnconfigure(j, minsize=col_widths[j], weight=1)
        
        # Update canvas scroll region to accommodate wider content
        if hasattr(self, 'canvas'):
            self.root.after(100, lambda: self.canvas.configure(scrollregion=self.canvas.bbox("all")))

    def update_weight_value(self, row_index):
        """Update weight values based on inapplicability checkboxes"""
        original_weights = [0.15, 0.12, 0.1, 0.08, 0.12, 0.08, 0.07, 0.1, 0.05, 0.06, 0.07]
        
        # Calculate redistribution
        total_checked_weight = sum(original_weights[i-1] for i in range(1, len(self.check_vars)) 
                                 if self.check_vars[i][3] and self.check_vars[i][3].get())
        
        unchecked_count = sum(1 for i in range(1, len(self.check_vars)) 
                            if not (self.check_vars[i][3] and self.check_vars[i][3].get()))
        
        # Update all weights
        for i in range(1, len(self.check_vars)):
            weight_cell = self.cells2[i][2]
            weight_cell.configure(state='normal')
            
            if self.check_vars[i][3] and self.check_vars[i][3].get():
                weight_cell.delete(0, tk.END)
                weight_cell.insert(0, "0")
            else:
                original_weight = original_weights[i-1]
                redistributed_weight = total_checked_weight / unchecked_count if unchecked_count > 0 else 0
                new_weight = original_weight + redistributed_weight
                weight_cell.delete(0, tk.END)
                weight_cell.insert(0, f"{new_weight:.3f}")
            
            weight_cell.configure(state='readonly')
        
        self.update_total_score()

    def update_total_score(self, event=None):
        """Calculate and update total score and risk level"""
        total_score = 0.0
        
        for i in range(1, len(self.combo_vars)):
            combo_var = self.combo_vars[i][1]
            if combo_var and combo_var.get():
                try:
                    value = float(combo_var.get())
                    weight_str = self.cells2[i][2].get()
                    if weight_str:
                        weight = float(weight_str)
                        total_score += (value - 1) * weight / 3
                except ValueError:
                    continue
        
        # Cap total score at 1.0 if it exceeds 0.99
        if total_score > 0.99:
            total_score = 1.0
        
        # Update total score cell
        total_cell = self.cells3[1][0]
        total_cell.configure(state='normal')
        total_cell.delete(0, tk.END)
        total_cell.insert(0, f"{total_score:.3f}")
        total_cell.configure(state='readonly')
        
        # Update risk level
        self.update_risk_level(total_score)

    def update_risk_level(self, total_score):
        """Update risk level based on total score"""
        if total_score <= 0.09:
            risk_level = "Very Low"
        elif total_score <= 0.39:
            risk_level = "Low"
        elif total_score <= 0.69:
            risk_level = "Medium"
        elif total_score <= 0.89:
            risk_level = "High"
        else:
            risk_level = "Very High"
        
        # Update risk level cell with colors
        risk_cell = self.cells3[1][1]        
        risk_cell.configure(state='normal')
        risk_cell.delete(0, tk.END)
        risk_cell.insert(0, risk_level)
        
        bg_color, fg_color = self.RISK_COLORS[risk_level]
        risk_cell.configure(readonlybackground=bg_color, fg=fg_color, insertbackground=fg_color)
        risk_cell.configure(state='readonly')

    def save_data(self):
        """Save data with visual feedback"""
        try:
            original_color = self.save_button.cget("bg")
            self.save_button.config(bg=self.COLORS['success'], text="💾 Saving...")
            self.root.update()
            
            self.root.after(500, lambda: self._complete_save(original_color))
        except Exception as e:
            messagebox.showerror("Error", f"Error during saving: {str(e)}")
            self.save_button.config(bg=self.COLORS['blue'], text="💾 Save Data")

    def _complete_save(self, original_color):
        """Complete save operation"""
        try:
            if DOCX_AVAILABLE:
                self._save_to_word()
            else:
                self._save_to_csv()
            
            self.save_button.config(bg=self.COLORS['success'], text="✅ Saved!")
            self.root.after(1000, lambda: self.save_button.config(bg=original_color, text="💾 Save Data"))
            
            messagebox.showinfo("Success", "Data saved successfully!\n\nBID Assessment report has been exported to Word.")
        except Exception as e:
            messagebox.showerror("Error", f"Error during saving: {str(e)}")
            self.save_button.config(bg=original_color, text="💾 Save Data")

    def _save_to_word(self):
        """Save to Word with formatting following Risk_Assessment_Optimized style"""
        # Automatic file name with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"BID_Assessment_{timestamp}.docx"
        
        # Create Output directory if it doesn't exist
        output_dir = os.path.join(get_base_path(), "Output")
        os.makedirs(output_dir, exist_ok=True)
        
        # Destination folder (Output directory)
        filepath = os.path.join(output_dir, filename)
        
        # Create Word document
        doc = Document()
        
        # Add content
        self.add_word_title(doc)
        self.add_risk_assessment_table(doc)
        self.add_results_table(doc)
        self.add_score_matrix_table(doc)
        
        # Save document
        doc.save(filepath)
        print(f"File saved as: {filepath}")

    def add_word_title(self, doc):
        """Adds the title to the Word document"""
        # Main title
        title = doc.add_heading('BID Phase - Risk Assessment', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Empty space

    def add_score_matrix_table(self, doc):
        """Adds the Score Matrix table to the Word document"""
        doc.add_heading('Score Matrix', level=1)
        
        # Create table with headers (5 columns, removed Weight)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        header_cells = table.rows[0].cells
        headers = ["Category", "Score 1 (Low)", "Score 2 (Significative)", "Score 3 (Moderate)", "Score 4 (High)"]
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Add data rows
        for i in range(1, len(self.TABLE1_DATA)):
            row_cells = table.add_row().cells
            for j, cell_text in enumerate(self.TABLE1_DATA[i]):
                row_cells[j].text = cell_text
                # Bold first column (categories)
                if j == 0:
                    row_cells[j].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()  # Space after table

    def add_risk_assessment_table(self, doc):
        """Adds the Risk Assessment table to the Word document"""
        doc.add_heading('Risk Assessment', level=1)
        
        # Create table with headers (4 columns, including Weight)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        header_cells = table.rows[0].cells
        headers = ["Category", "Value (1-4)", "Weight", "Inapplicability"]
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Add data rows
        for i in range(1, len(self.cells2)):
            row_cells = table.add_row().cells
            
            # Category name
            row_cells[0].text = self.cells2[i][0].get()
            row_cells[0].paragraphs[0].runs[0].bold = True
            
            # Value (ComboBox)
            combo_var = self.combo_vars[i][1]
            if combo_var:
                row_cells[1].text = combo_var.get()
            
            # Weight
            row_cells[2].text = self.cells2[i][2].get()
            
            # Inapplicability (Checkbox)
            check_var = self.check_vars[i][3]  # Back to index 3
            if check_var:
                row_cells[3].text = "✓" if check_var.get() else ""
        
        doc.add_paragraph()  # Space after table

    def add_results_table(self, doc):
        """Adds the Results table to the Word document"""
        doc.add_heading('Results', level=1)
        
        # Create table with headers
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        header_cells = table.rows[0].cells
        headers = ["Total Score", "Risk Level"]
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Add calculated values row
        row_cells = table.add_row().cells
        row_cells[0].text = self.cells3[1][0].get()
        row_cells[1].text = self.cells3[1][1].get()
        
        # Bold the values
        for cell in row_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Add reference table
        doc.add_paragraph()  # Space
        doc.add_paragraph("Risk Level Reference:")
        
        ref_table = doc.add_table(rows=1, cols=2)
        ref_table.style = 'Table Grid'
        ref_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Reference header
        ref_header_cells = ref_table.rows[0].cells
        ref_header_cells[0].text = "Score Range"
        ref_header_cells[1].text = "Level"
        
        for cell in ref_header_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Add reference data (skip first two rows which are calculated values)
        for i in range(3, len(self.cells3)):
            ref_row_cells = ref_table.add_row().cells
            ref_row_cells[0].text = self.cells3[i][0].get()
            ref_row_cells[1].text = self.cells3[i][1].get()
        
        doc.add_paragraph()  # Space after table

    def _save_to_csv(self):
        """Fallback CSV save when Word is not available"""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Error", "python-docx library not available!\nInstall with: pip install python-docx\n\nFalling back to CSV export.")
        
        import csv
        import os
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Save Risk Assessment Data
        risk_filename = f"BID_Risk_Assessment_{timestamp}.csv"
        
        # Create Output directory if it doesn't exist
        output_dir = os.path.join(get_base_path(), "Output")
        os.makedirs(output_dir, exist_ok=True)
        
        risk_filepath = os.path.join(output_dir, risk_filename)
        
        with open(risk_filepath, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            
            # Header (including Weight column)
            writer.writerow(["Category", "Value (1-4)", "Weight", "Inapplicability"])
            
            # Data rows
            for i in range(1, len(self.cells2)):
                row_data = []
                for j in range(4):  # Back to 4 columns
                    if j == 0:  # Category
                        cell_value = self.cells2[i][j].get()
                    elif j == 1:  # ComboBox value
                        combo_var = self.combo_vars[i][j]
                        cell_value = combo_var.get() if combo_var else ""
                    elif j == 2:  # Weight
                        cell_value = self.cells2[i][j].get()
                    elif j == 3:  # Checkbox (back to j == 3)
                        check_var = self.check_vars[i][j]
                        cell_value = "Yes" if (check_var and check_var.get()) else "No"
                    
                    row_data.append(cell_value)
                writer.writerow(row_data)
        
        # Save Results
        results_filename = f"BID_Results_{timestamp}.csv"
        results_filepath = os.path.join(output_dir, results_filename)
        
        with open(results_filepath, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            
            for i in range(len(self.cells3)):
                row_data = []
                for j in range(len(self.cells3[i])):
                    cell_value = self.cells3[i][j].get()
                    row_data.append(cell_value)
                writer.writerow(row_data)
        
        print(f"CSV files saved: {risk_filepath}, {results_filepath}")

    def show_help(self):
        """Show help window with criteria descriptions"""
        help_window = tk.Toplevel(self.root)
        help_window.title("Criteria Descriptions - Help")
        help_window.geometry("1200x700")  # Increased height to accommodate tool explanation
        help_window.configure(bg=self.COLORS['white'])
        help_window.resizable(True, True)
        
        # Center the window
        help_window.transient(self.root)
        help_window.grab_set()
        
        # Title
        title_label = tk.Label(help_window, text="Cybersecurity Criteria Descriptions", 
                              font=('Segoe UI', self.scaled_title_font, 'bold'),
                              bg=self.COLORS['white'], fg=self.COLORS['dark'])
        title_label.pack(pady=(20, 15))
        
        # Create scrollable frame for the table
        canvas = tk.Canvas(help_window, bg=self.COLORS['white'], highlightthickness=0)
        scrollbar = tk.Scrollbar(help_window, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=self.COLORS['white'])
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Create table frame
        table_frame = tk.Frame(scrollable_frame, bg=self.COLORS['white'])
        table_frame.pack(fill='both', expand=True, padx=15, pady=10)  # Reduced horizontal padding
        
        # Table headers
        header_frame = tk.Frame(table_frame, bg=self.COLORS['dark'], relief='ridge', bd=1)
        header_frame.pack(fill='x', pady=(0, 2))
        
        # Configure grid for better column control
        header_frame.grid_columnconfigure(0, weight=0, minsize=350)  # Increased width for criterion
        header_frame.grid_columnconfigure(1, weight=1)  # Flexible width for description
        
        criterion_header = tk.Label(header_frame, text="Criterion", font=('Segoe UI', self.scaled_font_size, 'bold'),
                                   bg=self.COLORS['dark'], fg=self.COLORS['white'], anchor='w',
                                   padx=15, pady=10)
        criterion_header.grid(row=0, column=0, sticky='ew')
        
        desc_header = tk.Label(header_frame, text="Description", font=('Segoe UI', self.scaled_font_size, 'bold'),
                              bg=self.COLORS['dark'], fg=self.COLORS['white'], anchor='w',
                              padx=15, pady=10)
        desc_header.grid(row=0, column=1, sticky='ew')
        
        # Add criteria rows
        for i, (criterion, description) in enumerate(self.CRITERIA_DESCRIPTIONS.items()):
            # Row frame
            row_color = self.COLORS['light'] if i % 2 == 0 else self.COLORS['white']
            row_frame = tk.Frame(table_frame, bg=row_color, relief='ridge', bd=1)
            row_frame.pack(fill='x', pady=1)
            
            # Configure grid for consistent column widths
            row_frame.grid_columnconfigure(0, weight=0, minsize=350)  # Increased width for criterion
            row_frame.grid_columnconfigure(1, weight=1)  # Flexible width for description
            
            # Criterion name (left column)
            criterion_label = tk.Label(row_frame, text=criterion, 
                                      font=('Segoe UI', self.scaled_font_size - 1, 'bold'),
                                      bg=row_color, fg=self.COLORS['dark'], anchor='nw',
                                      padx=15, pady=8, wraplength=320, justify='left')  # Increased wraplength
            criterion_label.grid(row=0, column=0, sticky='new')
            
            # Description (right column)
            desc_label = tk.Label(row_frame, text=description,
                                 font=('Segoe UI', self.scaled_font_size - 1),
                                 bg=row_color, fg='#495057', anchor='nw',
                                 padx=15, pady=8, wraplength=950, justify='left')  # Much larger wraplength
            desc_label.grid(row=0, column=1, sticky='new')
        
        # Add separator and tool explanation section
        separator_frame = tk.Frame(scrollable_frame, bg=self.COLORS['gray'], height=2)
        separator_frame.pack(fill='x', pady=(20, 15), padx=15)
        
        # Tool explanation title
        explanation_title = tk.Label(scrollable_frame, text="How the BID Risk Assessment Tool Works", 
                                    font=('Segoe UI', self.scaled_font_size + 2, 'bold'),
                                    bg=self.COLORS['white'], fg=self.COLORS['primary'])
        explanation_title.pack(pady=(10, 15), padx=15, anchor='w')
        
        # Tool explanation content
        explanation_frame = tk.Frame(scrollable_frame, bg=self.COLORS['light'], relief='ridge', bd=1)
        explanation_frame.pack(fill='x', padx=15, pady=(0, 20))
        
        explanation_text = """The BID (Bid Phase) Risk Assessment Tool helps evaluate cybersecurity risks during the bidding phase of space projects. Here's how to use it:

1. EVALUATION PROCESS:
   • For each of the 11 cybersecurity criteria, assess the quality of requirements in the ITT (Invitation to Tender)
   • Rate each criterion from 1 to 4 based on how well it is defined in the tender documents:
     - Score 1 (Low Risk): Well-defined, detailed requirements with clear specifications
     - Score 2 (Significative Risk): Partially defined with some details but lacking specificity
     - Score 3 (Moderate Risk): Vague or unclear requirements with minimal guidance
     - Score 4 (High Risk): No mention or very poor definition of the requirement

2. WEIGHTING SYSTEM:
   • Each criterion has a predefined weight based on its importance to overall cybersecurity
   • Higher weights are assigned to more critical security aspects like:
     - Cybersecurity Requirements (15%)
     - Security Architecture Constraints (12%)
     - Supply Chain Security (12%)

3. INAPPLICABILITY OPTION:
   • If a criterion is not applicable to the specific project, check the "Inapplicability" box
   • The tool will automatically redistribute that criterion's weight among the remaining applicable criteria
   • This ensures the assessment remains accurate for the project's specific context

4. AUTOMATIC CALCULATION:
   • The tool calculates a total risk score using the formula: Σ((Score-1) × Weight / 3)
   • This normalizes scores to a 0-1 scale where:
     - 0.0-0.1: Very Low Risk
     - 0.1-0.4: Low Risk
     - 0.4-0.7: Medium Risk
     - 0.7-0.9: High Risk
     - 0.9-1.0: Very High Risk

5. RESULTS INTERPRETATION:
   • Higher scores indicate greater cybersecurity risk due to poorly defined requirements
   • Use this assessment to:
     - Identify areas requiring clarification during the bidding process
     - Estimate additional effort needed for cybersecurity implementation
     - Make informed go/no-go decisions on tender participation
     - Plan appropriate cybersecurity measures and resources

6. EXPORT FUNCTIONALITY:
   • Save your assessment as a Word document for documentation and reporting
   • The export includes all evaluation details, scores, and risk analysis results"""
        
        explanation_label = tk.Label(explanation_frame, text=explanation_text,
                                   font=('Segoe UI', self.scaled_font_size - 1),
                                   bg=self.COLORS['light'], fg='#495057', anchor='nw',
                                   padx=20, pady=15, wraplength=1100, justify='left')
        explanation_label.pack(fill='both', expand=True)
        
        # Pack canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True, padx=(0, 5))
        scrollbar.pack(side="right", fill="y")
        
        # Mouse wheel scrolling for help window only
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        # Keep track of bound widgets for cleanup
        bound_widgets = []
        
        # Bind mouse wheel only to the help window and its children
        def bind_mousewheel(widget):
            widget.bind("<MouseWheel>", _on_mousewheel)
            bound_widgets.append(widget)
            for child in widget.winfo_children():
                bind_mousewheel(child)
        
        # Only bind to canvas and scrollable_frame to avoid conflicts
        canvas.bind("<MouseWheel>", _on_mousewheel)
        scrollable_frame.bind("<MouseWheel>", _on_mousewheel)
        bound_widgets.extend([canvas, scrollable_frame])
        
        # Ensure proper cleanup when window is closed
        def on_help_window_close():
            # Remove all mouse wheel bindings
            for widget in bound_widgets:
                try:
                    widget.unbind("<MouseWheel>")
                except:
                    pass  # Widget might be already destroyed
            help_window.destroy()
        
        help_window.protocol("WM_DELETE_WINDOW", on_help_window_close)
        
        # Focus on help window
        help_window.focus_set()

if __name__ == "__main__":
    root = tk.Tk()
    app = BIDOptimized(root)
    root.mainloop()
