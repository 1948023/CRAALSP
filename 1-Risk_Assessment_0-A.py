import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import math
import os
import sys
import csv
from datetime import datetime

# Configure environment for UTF-8 compatibility
os.environ['PYTHONIOENCODING'] = 'utf-8'

def get_base_path():
    """Get the base path for the application (works with both .py and .exe)"""
    if getattr(sys, 'frozen', False):
        # Running as compiled executable
        return os.path.dirname(sys.executable)
    else:
        # Running as script
        return os.path.dirname(os.path.abspath(__file__))

# Import for Word export/import
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class RiskAssessmentTool:
    """Optimized Risk Assessment Tool for space missions"""
      # Color configuration
    COLORS = {
        'primary': '#4a90c2', 'success': '#28a745', 'white': '#ffffff',
        'light': '#f8f9fa', 'dark': '#2c3e50', 'gray': '#6c757d',
        'criteria_header': '#5a67d8', 'criteria_bg': '#edf2f7',
        'asset_header': '#38b2ac', 'asset_bg': '#f0fff4'
    }    
    # Main table data
    THREATS = [
        "Data Corruption", "Physical/Logical Attack", "Interception/Eavesdropping",
        "Jamming", "Denial-of-Service", "Masquerade/Spoofing", "Replay",
        "Software Threats", "Unauthorized Access/Hijacking", 
        "Tainted hardware components", "Supply Chain"    ]
    
    def load_asset_categories_from_csv(self):
        """Load asset categories from Asset.csv (only categories and subcategories, no duplicates)"""
        assets_file = os.path.join(get_base_path(), "Asset.csv")
        asset_categories = []
        seen_combinations = set()
        
        try:
            with open(assets_file, 'r', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile, delimiter=';')
                for row in reader:
                    category = row.get('categories', '').strip()
                    subcategory = row.get('subCategories', '').strip()
                    
                    # Create tuple for category-subcategory combination
                    combination = (category, subcategory)
                    
                    # Add only if not already seen (avoid duplicates)
                    if combination not in seen_combinations and category and subcategory:
                        seen_combinations.add(combination)
                        asset_categories.append(combination)
            
            #print(f"[OK] Loaded {len(asset_categories)} unique asset categories from {assets_file}")
            return asset_categories
            
        except FileNotFoundError:
            #print(f"[NO] File not found: {assets_file}")
            # Fallback asset categories
            return [
                ("Ground", "Ground Stations"), ("Ground", "Mission Control"),
                ("Ground", "Data Processing Centers"), ("Ground", "Remote Terminals"),
                ("Ground", "User Ground Segment"), ("Space", "Platform"),
                ("Space", "Payload"), ("Link", "Link"), ("User", "User")
            ]
        except Exception as e:
            #print(f"[NO] Error loading asset categories: {e}")
            return [
                ("Ground", "Ground Stations"), ("Ground", "Mission Control"),
                ("Ground", "Data Processing Centers"), ("Ground", "Remote Terminals"),
                ("Ground", "User Ground Segment"), ("Space", "Platform"),
                ("Space", "Payload"), ("Link", "Link"), ("User", "User")
            ]
    
    # Criteria table data (5x6 + header) - Transposed format
    CRITERIA_DATA = [
        ["Score", "Vulnerability Level", "Access Control", "Defense Capability", "Operational Impact", "Recovery Time"],
        ["Score 1 (Very Low)", "No know or already resolved vulnerabilities", "Access strongly protected by physical/logical measures and isolated environment", "Multi-level validated countermeasures with real-time automated detection", "No impact thanks to redundancy with predefined automated response", "Immediate restoration with automated procedures"],
        ["Score 2 (Low)", "Know vulnerability, mitigate throught hardening and patches", "Moderately protected access with some isolation controls", "Robust countermeasures with automated but decentralized detection", "Negligible impact, quick response and system easily restored", "Quick recovery within hours to days using standard procedures"],
        ["Score 3 (Moderate)", "Know vulnerability, but only partially mitigated", "Standard access protection with basic controls", "Limited countermeasures with manual detection only", "Medium impact with manual response, but mission continues", "Manual recovery requiring weeks of coordinated effort"],
        ["Score 4 (High)", "Known vulnerability, with no effective mitigation", "Access easily accessible by remote attackers", "Weak countermeasure with occasional detection", "Serious impact with slow response, mission temporarily interrupted", "Complex recovery requiring months of specialized intervention"],
        ["Score 5 (Very High)", "Actively exploitable vulnerability, with no defense", "Completely open or physically accessible access", "No countermeasures or detection capabilities", "Permanent loss of assets or mission with no response capability", "Impossible recovery or permanent system loss"]    ]
    
    # Risk matrix
    RISK_MATRIX = {
        ("Very High", "Very High"): "Very High", ("Very High", "High"): "Very High",
        ("Very High", "Medium"): "High", ("Very High", "Low"): "High",
        ("Very High", "Very Low"): "Medium", ("High", "Very High"): "Very High",
        ("High", "High"): "High", ("High", "Medium"): "High",
        ("High", "Low"): "Medium", ("High", "Very Low"): "Low",
        ("Medium", "Very High"): "High", ("Medium", "High"): "High",
        ("Medium", "Medium"): "Medium", ("Medium", "Low"): "Low",
        ("Medium", "Very Low"): "Low", ("Low", "Very High"): "Medium",
        ("Low", "High"): "Medium", ("Low", "Medium"): "Low",
        ("Low", "Low"): "Low", ("Low", "Very Low"): "Very Low",
        ("Very Low", "Very High"): "Low", ("Very Low", "High"): "Low",
        ("Very Low", "Medium"): "Low", ("Very Low", "Low"): "Very Low",
        ("Very Low", "Very Low"): "Very Low"    }
    # Security controls for each threat
    THREAT_COUNTERMEASURES = {
        "Data Corruption": [
            "Configuration Management", "Tamper resistant body", "Tamper Protection", 
            "Disable Physical Ports", "Anti-counterfeit Hardware", "Secure disposal or reuse of equipment",
            "Access-based network segmentation", "Vulnerability Management", "Malware Protection",
            "ASIC/FPGA Manufacturing"
        ],
        "Physical/Logical Attack": [
            "A tamper resistant body", "Satellite Unit RF Encryption", "Traffic Flow Security",
            "Power Masking", "Secure disposal or reuse of equipment", "Access-based network segmentation",
            "Information classification and labelling", "Vulnerability Management", "Malware Protection"
        ],
        "Interception/Eavesdropping": [
            "Communications Security", "Satellite Unit RF Encryption", "Traffic Flow Security",
            "Power Masking", "Access-based network segmentation", "Information classification and labelling"
        ],
        "Jamming": [
            "Resilient Position Navigation and Timing", "Communication Physical Medium Space-Based",
            "Radio Frequency Mapping", "Antenna Nulling and Adaptive Filtering",
            "Defensive Jamming and Spoofing", "Emergency power sources",
            "Real-time physics model-based system verification"
        ],
        "Denial-of-Service": [
            "Security of Power Systems", "System redundancy", "Incident Recovery Plan",
            "Emergency power sources", "Traffic Flow Security",
            "Critical Services Delivery Requirements"
        ],
        "Masquerade/Spoofing": [
            "OSAM Dual Authorization", "Multi factor authentication", "Relay Protection",
            "Smart Contracts", "Resilient Position Navigation and Timing"
        ],
        "Replay": [
            "Relay Protection", "Satellite Unit RF Encryption", "On-board Message Encryption",
            "Session Termination", "Real-time physics model-based system verification"
        ],
        "Software Threats": [
            "Coding Standard", "Malware Protection", "Vulnerability scanning", "Vulnerability Management",
            "Software Updates", "Dynamic Code Analysis", "Static Code Analysis", "Process ID whitelisting",
            "Software Bill of Materials"
        ],
        "Unauthorized Access/Hijacking": [
            "Access rights", "Identity management", "Remote access management", "Multi factor authentication",
            "Access-based network segmentation", "Backdoor Commands"
        ],
        "Tainted hardware components": [
            "Anti-counterfeit Hardware", "ASIC/FPGA Manufacturing", "Tamper Protection",
            "Supplier Security Management"
        ],
        "Supply Chain": [
            "Supplier Security Management", "Software Bill of Materials", "Software Supply Chain Integrity",
            "Outsourced development", "Cloud Cybersecurity Measures"
        ]    }
    
    # Available mission types
    MISSION_TYPES = [
        "Insert type of mission",
        "Earth Observation Mission",
        "Communication Satellite", 
        "Scientific Mission",
        "Navigation Satellite",
        "On-Orbit Service"
    ]
    def __init__(self, root):
        self.root = root
        self.root.title("Risk Assessment Tool - Phase 0-A")
        self.root.state('zoomed')        
        self.root.configure(bg=self.COLORS['white'])
        
        # Load asset categories from CSV
        self.ASSET_CATEGORIES = self.load_asset_categories_from_csv()
        
        # Data for threats and calculations
        self.threat_data = {}  # Saved data for threat
        self.combo_vars = {}   # ComboBox variables
        self.impact_entries = {}  # Table widgets
        
        # Variable for mission type
        self.mission_type_var = tk.StringVar(value=self.MISSION_TYPES[0])
        
        self.create_interface()
    
    def disable_mousewheel_on_combobox(self, combo):
        """Intelligently handle mouse wheel on combobox to prevent accidental value changes while allowing scroll"""
        def on_mousewheel(event):
            # Check if the combobox dropdown is open
            try:
                if combo.tk.call('ttk::combobox::PopdownIsVisible', combo):
                    # If dropdown is open, allow normal combobox behavior
                    return
                else:
                    # If dropdown is closed, prevent value changes but allow window scrolling
                    # Find the parent canvas to continue scrolling
                    widget = event.widget
                    # Walk up the widget hierarchy to find a canvas
                    while widget:
                        if isinstance(widget, tk.Canvas):
                            widget.yview_scroll(int(-1*(event.delta/120)), "units")
                            break
                        widget = widget.master
                    return "break"  # Prevent combobox value change
            except:
                # Fallback: prevent value changes but allow window scrolling
                widget = event.widget
                while widget:
                    if isinstance(widget, tk.Canvas):
                        widget.yview_scroll(int(-1*(event.delta/120)), "units")
                        break
                    widget = widget.master
                return "break"  # Prevent combobox value change
        
        combo.bind("<MouseWheel>", on_mousewheel)

    def setup_global_mousewheel(self, widget, canvas):
        """Setup global mouse wheel scrolling for any widget relative to a canvas"""
        def on_global_mousewheel(event):
            # Scroll the canvas when mouse wheel is used anywhere in the widget
            try:
                canvas.yview_scroll(int(-1*(event.delta/120)), "units")
            except:
                pass  # Ignore errors if canvas is not scrollable
        
        # Bind to the widget and all its children recursively
        def bind_mousewheel_recursive(w):
            # Always bind to non-combobox widgets
            if not isinstance(w, ttk.Combobox):
                w.bind("<MouseWheel>", on_global_mousewheel)
            
            # Recursively bind to all children
            for child in w.winfo_children():
                bind_mousewheel_recursive(child)
        
        # Start recursive binding
        bind_mousewheel_recursive(widget)
        
        # Also bind directly to the main widget and canvas for safety
        widget.bind("<MouseWheel>", on_global_mousewheel)
        canvas.bind("<MouseWheel>", on_global_mousewheel)

    def ensure_mousewheel_on_table_cells(self):
        """Ensure all threat table cells have mouse wheel scrolling - for non-scrollable tables"""
        def on_cell_mousewheel(event):
            # For non-scrollable main tables, we can add a gentle visual feedback
            # or just ignore since there's no scroll needed
            pass
        
        # Apply to all threat cells in main table
        for threat, cell in self.threat_cells.items():
            # For main table without scroll, we don't need special handling
            # but we can still bind for consistency
            pass

    def setup_asset_table_mousewheel(self, canvas):
        """Setup mouse wheel scrolling for asset table in threat analysis window"""
        def on_asset_table_mousewheel(event):
            try:
                canvas.yview_scroll(int(-1*(event.delta/120)), "units")
            except:
                pass
        
        # Apply to all combo boxes and cells in asset assessment table
        for asset_key in self.combo_vars:
            for combo_idx, combo_var in self.combo_vars[asset_key].items():
                # We don't bind directly to comboboxes since they have their own handler
                pass
        
        # Apply to calculated cells
        for asset_key in self.impact_entries:
            for col_idx, widget in self.impact_entries[asset_key].items():
                if hasattr(widget, 'bind') and not isinstance(widget, ttk.Combobox):
                    widget.bind("<MouseWheel>", on_asset_table_mousewheel)
    
    def create_interface(self):
        """Creates the main interface"""
        # Header
        header = tk.Frame(self.root, bg=self.COLORS['light'], height=60)
        header.pack(fill='x')
        header.pack_propagate(False)
        
        tk.Label(header, text="Risk Assessment Tool - Phase 0-A", 
                font=('Segoe UI', 16, 'bold'),
                bg=self.COLORS['light'], fg=self.COLORS['dark']).pack(pady=15)
        
        # Container principale
        main_frame = tk.Frame(self.root, bg=self.COLORS['white'])
        main_frame.pack(fill='both', expand=True, padx=20, pady=20)
        
        # Threat table
        self.create_threats_table(main_frame)
        # Buttons
        self.create_buttons(main_frame)
    
    def create_threats_table(self, parent):
        """Creates the threat table"""
        # Main container
        main_container = tk.Frame(parent, bg=self.COLORS['white'])
        main_container.pack(fill='both', expand=True, pady=(0, 20))

        # Mission type selector (separate from the table)
        mission_frame = tk.LabelFrame(main_container, text="Mission Configuration",
                                     font=('Segoe UI', 11, 'bold'),
                                     bg=self.COLORS['white'], fg=self.COLORS['primary'],
                                     padx=15, pady=10)
        mission_frame.pack(fill='x', pady=(0, 10))
        
        tk.Label(mission_frame, text="Mission Type:",
                font=('Segoe UI', 10, 'bold'),
                bg=self.COLORS['white'], fg=self.COLORS['dark']).pack(anchor='w')
        
        mission_combo = ttk.Combobox(mission_frame,
                                   textvariable=self.mission_type_var,
                                   values=self.MISSION_TYPES,
                                   font=('Segoe UI', 10),
                                   state='readonly')
        mission_combo.pack(fill='x', pady=(5, 0))
        
        # Disable mouse wheel on mission combobox
        self.disable_mousewheel_on_combobox(mission_combo)

        # Threat table (separate from the mission selector)
        table_frame = tk.LabelFrame(main_container, text="Threat Risk Levels",
                                   font=('Segoe UI', 12, 'bold'),
                                   bg=self.COLORS['white'], fg=self.COLORS['primary'],
                                   padx=20, pady=15)
        table_frame.pack(fill='both', expand=True)
        
        # Header
        headers = ["Threat", "Risk Level"]
        for j, header in enumerate(headers):
            cell = tk.Label(table_frame, text=header,
                           font=('Segoe UI', 11, 'bold'),
                           bg=self.COLORS['primary'], fg=self.COLORS['white'],
                           relief='ridge', bd=1)
            cell.grid(row=0, column=j, sticky='ew', padx=1, pady=1, ipady=8)
        
        # Data Rows
        self.threat_cells = {}
        for i, threat in enumerate(self.THREATS, 1):
            # Threat name
            name_cell = tk.Label(table_frame, text=threat,
                               font=('Segoe UI', 10),
                               bg=self.COLORS['white'], fg=self.COLORS['dark'],
                               relief='ridge', bd=1, anchor='w')
            name_cell.grid(row=i, column=0, sticky='ew', padx=1, pady=1, ipady=5)
            
            # Risk level
            risk_cell = tk.Label(table_frame, text="",
                               font=('Segoe UI', 10),
                               bg=self.COLORS['white'], fg=self.COLORS['dark'],
                               relief='ridge', bd=1)
            risk_cell.grid(row=i, column=1, sticky='ew', padx=1, pady=1, ipady=5)
            
            self.threat_cells[threat] = risk_cell

        # Grid configuration
        for j in range(2):
            table_frame.grid_columnconfigure(j, weight=1)
    
    def create_buttons(self, parent):
        """Creates the buttons"""
        button_frame = tk.Frame(parent, bg=self.COLORS['white'])
        button_frame.pack(fill='x', pady=20)

        # Main container with three sections
        main_container = tk.Frame(button_frame, bg=self.COLORS['white'])
        main_container.pack()

        # Left section - Export/Import 0-A buttons
        left_frame = tk.Frame(main_container, bg=self.COLORS['white'])
        left_frame.pack(side='left', padx=(0, 30))
        
        if DOCX_AVAILABLE:
            export_word_btn = tk.Button(left_frame, text="EXPORT 0-A REPORT",
                                       font=('Segoe UI', 10, 'bold'),
                                       bg="#0E831E", fg=self.COLORS['white'],
                                       relief='flat', padx=20, pady=2,
                                       command=self.export_to_word)
            export_word_btn.pack(pady=(0, 2))
            
            import_word_btn = tk.Button(left_frame, text="IMPORT 0-A REPORT",
                                       font=('Segoe UI', 10, 'bold'),
                                       bg='#8e44ad', fg=self.COLORS['white'],
                                       relief='flat', padx=19.5, pady=2,
                                       command=self.import_from_word)
            import_word_btn.pack(pady=(2, 0))
        else:
            no_docx_label = tk.Label(left_frame, text="Word export/import unavailable\ninstall python-docx",
                                   font=('Segoe UI', 9),
                                   bg=self.COLORS['white'], fg=self.COLORS['gray'],
                                   justify='center')
            no_docx_label.pack()

        # Center section - THREAT ANALYSIS button
        center_frame = tk.Frame(main_container, bg=self.COLORS['white'])
        center_frame.pack(side='left', padx=30)
        
        threat_btn = tk.Button(center_frame, text="THREAT ANALYSIS",
                              font=('Segoe UI', 14, 'bold'),
                              bg=self.COLORS['primary'], fg=self.COLORS['white'],
                              relief='flat', padx=40, pady=15,
                              command=self.open_threat_window)
        threat_btn.pack()

        # Right section - Import Legacy Report button
        right_frame = tk.Frame(main_container, bg=self.COLORS['white'])
        right_frame.pack(side='left', padx=(30, 0))
        
        import_legacy_btn = tk.Button(right_frame, text="IMPORT MISSION\nANALYSIS REPORT",
                                     font=('Segoe UI', 11, 'bold'),
                                     bg='#e67e22', fg=self.COLORS['white'],
                                     relief='flat', padx=20, pady=12,
                                     command=self.import_legacy_report,
                                     justify='center')
        import_legacy_btn.pack()
    
    def open_threat_window(self):
        """Open Threat Analysis window"""
        window = tk.Toplevel(self.root)
        window.title("Threat Analysis")
        window.geometry("1400x800")
        window.configure(bg=self.COLORS['white'])
        window.transient(self.root)
        window.grab_set()
        
        # Header
        header = tk.Frame(window, bg=self.COLORS['light'], height=50)
        header.pack(fill='x')
        header.pack_propagate(False)
        
        tk.Label(header, text="Threat Analysis",
                font=('Segoe UI', 14, 'bold'),
                bg=self.COLORS['light'], fg=self.COLORS['dark']).pack(pady=12)

        # Main container with scroll
        self.create_threat_content(window)
    
    def create_threat_content(self, window):
        """Creates the threat content window"""
        # Scrollable canvas
        canvas = tk.Canvas(window, bg=self.COLORS['white'])
        scrollbar = tk.Scrollbar(window, orient="vertical", command=canvas.yview)
        content_frame = tk.Frame(canvas, bg=self.COLORS['white'])
        
        content_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=content_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True, padx=20, pady=20)
        scrollbar.pack(side="right", fill="y")

        # Criteria table
        self.create_criteria_table(content_frame)
        
        # Threat selection
        threat_frame = tk.Frame(content_frame, bg=self.COLORS['white'])
        threat_frame.pack(fill='x', pady=(20, 20))
        
        tk.Label(threat_frame, text="Select Threat:",
                font=('Segoe UI', 11, 'bold'),
                bg=self.COLORS['white'], fg=self.COLORS['dark']).pack(anchor='w')
        
        self.selected_threat_var = tk.StringVar()
        threat_combo = ttk.Combobox(threat_frame,
                                   textvariable=self.selected_threat_var,
                                   values=self.THREATS,
                                   font=('Segoe UI', 10),
                                   state='readonly')
        threat_combo.pack(fill='x', pady=(5, 0))
        threat_combo.bind('<<ComboboxSelected>>', self.load_threat_data)
        
        # Disable mouse wheel on threat combobox
        self.disable_mousewheel_on_combobox(threat_combo)
        
        # Asset table
        self.create_asset_table(content_frame)

        # Setup mouse wheel for asset table elements
        self.setup_asset_table_mousewheel(canvas)

        # Buttons frame
        buttons_frame = tk.Frame(content_frame, bg=self.COLORS['white'])
        buttons_frame.pack(pady=20)
        
        # Save button
        save_btn = tk.Button(buttons_frame, text="SAVE ASSESSMENT",
                            font=('Segoe UI', 11, 'bold'),
                            bg=self.COLORS['success'], fg=self.COLORS['white'],
                            relief='flat', padx=25, pady=10,
                            command=lambda: self.save_threat_assessment(window))
        save_btn.pack(side='left', padx=(0, 10))
        
        # Help button
        help_btn = tk.Button(buttons_frame, text="❓ Help",
                            font=('Segoe UI', 11, 'bold'),
                            bg=self.COLORS['gray'], fg=self.COLORS['white'],
                            relief='flat', padx=20, pady=10,
                            command=self.show_help)
        help_btn.pack(side='left')

        # Setup global mouse wheel scrolling for the entire content frame
        self.setup_global_mousewheel(content_frame, canvas)
            
    def create_criteria_table(self, parent):
        """Creates the assessment criteria table"""
        criteria_container = tk.LabelFrame(parent, 
                                         text="Assessment Criteria",
                                         font=('Segoe UI', 12, 'bold'),
                                         bg=self.COLORS['white'], 
                                         fg=self.COLORS['primary'], 
                                         padx=20, 
                                         pady=15,                                         
                                         relief='ridge', 
                                         bd=2)
        criteria_container.pack(fill='x', pady=(0, 20))

        # Creates the cells of the criteria table
        for i, row in enumerate(self.CRITERIA_DATA):
            for j, cell_text in enumerate(row):                    
                if i == 0:  # Header row                    
                    cell = tk.Label(criteria_container, text=cell_text,
                                   font=('Segoe UI', 10, 'bold'),
                                   bg=self.COLORS['criteria_header'], 
                                   fg=self.COLORS['white'],
                                   relief='ridge', 
                                   bd=1,                                   
                                   anchor='center',
                                   justify='center',
                                   wraplength=180,  
                                   width=22,        
                                   height=3,
                                   padx=3,         
                                   pady=2)
                else:  # Data rows
                    
                    font_weight = 'bold' if j == 0 else 'normal'                        
                    cell = tk.Label(criteria_container, text=cell_text,
                                   font=('Segoe UI', 9, font_weight),
                                   bg=self.COLORS['criteria_bg'],
                                   fg=self.COLORS['dark'],
                                   relief='ridge',
                                   bd=1,
                                   anchor='nw',
                                   justify='left',
                                   wraplength=180,  
                                   width=22,        
                                   height=4,        
                                   padx=6,          
                                   pady=3)          
                
                cell.grid(row=i, column=j, padx=2, pady=2, sticky='ew', ipady=5)

        # Grid configuration with uniform column sizes
        for j in range(6):
            criteria_container.grid_columnconfigure(j, weight=1, minsize=180, uniform="criteria_cols")  
        
        for i in range(6):  # 5 data rows + 1 header
            criteria_container.grid_rowconfigure(i, minsize=60, uniform="criteria_rows")  
    
    def create_asset_table(self, parent):
        """Creates the asset assessment table"""
        table_frame = tk.LabelFrame(parent, text="Asset Assessment (Values 1-5)",
                                   font=('Segoe UI', 11, 'bold'),
                                   bg=self.COLORS['white'], fg=self.COLORS['primary'],
                                   padx=15, pady=15)        
        table_frame.pack(fill='both', expand=True)
        
        # Headers
        headers = ["Category", "Sub-Category", "Vulnerability", "Access", "Defense", 
                  "Operational Impact", "Recovery", "Likelihood", "Impact", "Risk"]
        
        for j, header in enumerate(headers):
            cell = tk.Label(table_frame, text=header,
                           font=('Segoe UI', 10, 'bold'),
                           bg=self.COLORS['primary'], fg=self.COLORS['white'],
                           relief='ridge', bd=1, width=12)
            cell.grid(row=0, column=j, padx=1, pady=1, sticky='ew')
        
        # Reset storage
        self.impact_entries = {}
        self.combo_vars = {}
        
        # Assessment Rows - Dynamic based on loaded asset categories
        for i in range(len(self.ASSET_CATEGORIES)):
            category, subcategory = self.ASSET_CATEGORIES[i]
            asset_key = f"{i+1}_probability"  # Unique key for asset
            # Category (read-only)
            cat_cell = tk.Label(table_frame, text=category,
                               font=('Segoe UI', 9, 'bold'),
                               bg=self.COLORS['light'], fg=self.COLORS['dark'],
                               relief='ridge', bd=1, width=10)
            cat_cell.grid(row=i+1, column=0, padx=1, pady=1, sticky='ew')
            
            # Sub-Category (read-only)
            sub_cell = tk.Label(table_frame, text=subcategory,
                               font=('Segoe UI', 9),
                               bg=self.COLORS['light'], fg=self.COLORS['dark'],
                               relief='ridge', bd=1, width=15)
            sub_cell.grid(row=i+1, column=1, padx=1, pady=1, sticky='ew')

            # Storage for this row
            row_entries = {}
            self.combo_vars[asset_key] = {}
            
            # Writable columns (2-6: Vulnerability, Access, Defense, Operational Impact, Recovery)
            for j in range(2, 7):
                combo_var = tk.StringVar(value="")                
                combo = ttk.Combobox(table_frame,
                                    textvariable=combo_var,
                                    values=["", "1", "2", "3", "4", "5"],
                                    font=('Segoe UI', 9),
                                    width=8, state='readonly')
                combo.grid(row=i+1, column=j, padx=1, pady=1, sticky='ew')
                
                # Disable mouse wheel on combobox
                self.disable_mousewheel_on_combobox(combo)
                
                row_entries[j-2] = combo  # 0-based index
                self.combo_vars[asset_key][j-2] = combo_var

                # Bind calculations
                if j <= 4:  # Vulnerability, Access, Defense -> Likelihood
                    combo_var.trace_add('write', lambda *args, key=asset_key: self.calculate_likelihood(key))
                elif j <= 6:  # Operational Impact, Recovery -> Impact
                    combo_var.trace_add('write', lambda *args, key=asset_key: self.calculate_impact(key))
            # Colonne calcolate (7-9: Likelihood, Impact, Risk) - read-only
            for j in range(7, 10):
                calc_cell = tk.Label(table_frame, text="",
                                   font=('Segoe UI', 9),
                                   bg=self.COLORS['light'], fg=self.COLORS['dark'],
                                   relief='ridge', bd=1, width=10)
                calc_cell.grid(row=i+1, column=j, padx=1, pady=1, sticky='ew')
                row_entries[j-2] = calc_cell
            
            self.impact_entries[asset_key] = row_entries        
        # Column 0 (Category):
        table_frame.grid_columnconfigure(0, weight=1, minsize=120, uniform="small_cols")
        # Column 1 (Sub-Category): 
        table_frame.grid_columnconfigure(1, weight=1, minsize=180, uniform="sub_category_col")
        # Columns 2-9:
        for j in range(2, 10):
            table_frame.grid_columnconfigure(j, weight=1, minsize=120, uniform="small_cols")
        
        for i in range(10):  # 9 data rows + 1 header
            table_frame.grid_rowconfigure(i, minsize=40, uniform="rows")
    
    def calculate_likelihood(self, key):
        """Calculates Likelihood using quadratic mean of three criteria"""
        if key not in self.combo_vars:
            return

        # Get values Vulnerability, Access, Defense (columns 0,1,2)
        values = []
        for col_idx in [0, 1, 2]:
            if col_idx not in self.combo_vars[key]:
                continue
                self.update_display(key, 5, "")  # Likelihood column
                return
            
            value_str = self.combo_vars[key][col_idx].get().strip()
            if not value_str or value_str == "0":
                continue
                self.update_display(key, 5, "")
                return
            
            try:
                values.append(float(value_str))
            except ValueError:
                self.update_display(key, 5, "")
                return
        
        #if len(values) != 3:
        #    self.update_display(key, 5, "")
        #    return

        # Calculate Likelihood using quadratic mean
        quadratic_mean = math.sqrt(sum(x**2 for x in values) / len(values))
        likelihood = (quadratic_mean - 1) / 4  # Normalize [1,5] -> [0,1]
        likelihood = max(0.0, min(1.0, likelihood))

        # Update display with category instead of numeric value
        likelihood_category = self.value_to_category(likelihood)
        self.update_display(key, 5, likelihood_category)  # Likelihood column

        # Recalculate risk if Impact is also available
        self.calculate_risk(key)

        # Update main table in real-time also for likelihood
        self.update_main_table_risk_realtime()

    def get_saved_likelihood(self, threat_name, asset_num):
        """Get saved likelihood for specific threat/asset using only base calculation"""
        if threat_name not in self.threat_data:
            return 0.0
        
        asset_key = f"{asset_num}_probability"
        if asset_key not in self.threat_data[threat_name]:
            return 0.0
        
        row_data = self.threat_data[threat_name][asset_key]
        
        try:
            # Check values for Vulnerability, Access, Defense
            if not isinstance(row_data, dict):
                return 0.0
            
            values = []
            for i in [0, 1, 2]:  # Vulnerability, Access, Defense
                if str(i) not in row_data:
                    return 0.0
                val = row_data[str(i)]
                if not val or val == "0":
                    return 0.0
                values.append(float(val))
            
            if len(values) == 3:
                # Calculate base likelihood using quadratic mean
                quadratic_mean = math.sqrt(sum(x**2 for x in values) / 3)
                likelihood = (quadratic_mean - 1) / 4  # Normalize [1,5] -> [0,1]
                return max(0.0, min(1.0, likelihood))
        
        except (ValueError, KeyError, TypeError):
            pass
        
        return 0.0
    
    def calculate_impact(self, key):
        """Calculates Impact as the quadratic mean of Operational Impact and Recovery Time"""
        if key not in self.combo_vars:
            return

        # Get Operational Impact, Recovery values (columns 3,4)
        values = []
        for col_idx in [3, 4]:
            if col_idx not in self.combo_vars[key]:
                continue
                self.update_display(key, 6, "")  # Impact column
                return
            
            value_str = self.combo_vars[key][col_idx].get().strip()
            if not value_str or value_str == "0":
                continue
                self.update_display(key, 6, "")
                return
            
            try:
                values.append(float(value_str))
            except ValueError:
                self.update_display(key, 6, "")
                return
        
        # Quadratic mean normalized
        quadratic_mean = math.sqrt(sum(x**2 for x in values) / len(values))
        impact = (quadratic_mean - 1) / 4  # [1,5] -> [0,1]
        impact = max(0.0, min(1.0, impact))

        # Update display with category instead of numeric value
        impact_category = self.value_to_category(impact)
        self.update_display(key, 6, impact_category)  # Impact column

        # Recalculate risk
        self.calculate_risk(key)

    def calculate_risk(self, key):
        """Calculates Risk using the Likelihood x Impact matrix"""
        if key not in self.impact_entries:
            return

        # Get Likelihood and Impact (now they are already categories)
        likelihood_widget = self.impact_entries[key][5]  # Likelihood column
        impact_widget = self.impact_entries[key][6]      # Impact column
        
        likelihood_cat = likelihood_widget.cget('text') if hasattr(likelihood_widget, 'cget') else ""
        impact_cat = impact_widget.cget('text') if hasattr(impact_widget, 'cget') else ""
        
        if not likelihood_cat or not impact_cat:
            self.update_display(key, 7, "")  # Risk column
            return

        # Check if they are valid categories
        valid_categories = ["Very Low", "Low", "Medium", "High", "Very High"]
        if likelihood_cat in valid_categories and impact_cat in valid_categories:
            # Get risk from matrix
            risk_level = self.RISK_MATRIX.get((likelihood_cat, impact_cat), "")
            self.update_display(key, 7, risk_level)  # Risk column

            # Update main table in real-time
            self.update_main_table_risk_realtime()
        else:
            self.update_display(key, 7, "")
    
    def value_to_category(self, value):
        """Converts numeric value to category"""
        if value <= 0.1:
            return "Very Low"
        elif value <= 0.4:
            return "Low"
        elif value <= 0.7:
            return "Medium"
        elif value <= 0.9:
            return "High"
        else:
            return "Very High"
    
    def update_display(self, key, col_index, value):
        """Updates the display of a cell"""
        if key in self.impact_entries and col_index in self.impact_entries[key]:
            widget = self.impact_entries[key][col_index]
            if hasattr(widget, 'config'):
                widget.config(text=value)
    
    def load_threat_data(self, event=None):
        """Loads data for selected threat"""
        selected_threat = self.selected_threat_var.get()

        # Clear all fields
        for key in self.impact_entries:
            self.clear_asset_data(key)

        # Load data if it exists
        if selected_threat and selected_threat in self.threat_data:
            threat_data = self.threat_data[selected_threat]
            
            for key, row_data in threat_data.items():
                if key in self.combo_vars:
                    for col_idx, value in row_data.items():
                        if int(col_idx) in self.combo_vars[key]:
                            self.combo_vars[key][int(col_idx)].set(value)

        # Recalculate everything
        for key in self.impact_entries:
            self.calculate_likelihood(key)
            self.calculate_impact(key)
    
    def clear_asset_data(self, key):
        """Clears data for an asset"""
        if key in self.combo_vars:
            for combo_var in self.combo_vars[key].values():
                combo_var.set("")
        
        if key in self.impact_entries:
            for col_idx in [5, 6, 7]:  # Likelihood, Impact, Risk
                if col_idx in self.impact_entries[key]:
                    self.update_display(key, col_idx, "")
    def save_threat_assessment(self, window):
        """Saves current threat assessment"""
        selected_threat = self.selected_threat_var.get()
        if not selected_threat:
            messagebox.showwarning("Warning", "Please select a threat first!")
            return

        # Collect data
        threat_data = {}
        for key in self.combo_vars:
            row_data = {}
            for col_idx, combo_var in self.combo_vars[key].items():
                value = combo_var.get().strip()
                if value:
                    row_data[str(col_idx)] = value
            
            if row_data:
                threat_data[key] = row_data
        
        # Save data
        self.threat_data[selected_threat] = threat_data

        # Update main table with maximum risks of ALL threats
        self.update_all_threats_in_main_table()
        
        messagebox.showinfo("Success", f"Assessment for '{selected_threat}' saved!")
        window.destroy()
    
    def update_main_table_risk(self, threat_name):
        """Updates main table with maximum risk"""
        if threat_name not in self.threat_data or threat_name not in self.threat_cells:
            return

        # Find maximum risk among all assets
        risk_priorities = {"Very High": 5, "High": 4, "Medium": 3, "Low": 2, "Very Low": 1, "": 0}
        max_risk = ""
        max_priority = 0

        # Check all assets
        for key in self.impact_entries:
            risk_widget = self.impact_entries[key].get(7)  # Risk column
            if risk_widget and hasattr(risk_widget, 'cget'):
                risk_text = risk_widget.cget('text')
                priority = risk_priorities.get(risk_text, 0)
                if priority > max_priority:
                    max_priority = priority
                    max_risk = risk_text        # Update main table
        self.threat_cells[threat_name].config(text=max_risk)
    
    def update_all_threats_in_main_table(self):
        """Updates main table with maximum risks of all saved threats"""
        risk_priorities = {"Very High": 5, "High": 4, "Medium": 3, "Low": 2, "Very Low": 1, "": 0}

        # For each threat that has saved data
        for threat_name in self.threat_data.keys():
            if threat_name not in self.threat_cells:
                continue
            
            max_risk = ""
            max_priority = 0

            # Calculate maximum risk for this threat
            threat_data = self.threat_data[threat_name]
            
            for asset_key, asset_data in threat_data.items():
                # Calculate likelihood for this asset
                likelihood = self.calculate_likelihood_from_saved_data(threat_name, asset_key, asset_data)

                # Calculate impact for this asset
                impact = self.calculate_impact_from_saved_data(asset_data)

                # Calculate risk if both are available
                if likelihood >= 0 and impact >= 0:
                    likelihood_cat = self.value_to_category(likelihood)
                    impact_cat = self.value_to_category(impact)
                    risk_level = self.RISK_MATRIX.get((likelihood_cat, impact_cat), "")
                    
                    priority = risk_priorities.get(risk_level, 0)
                    if priority > max_priority:
                        max_priority = priority
                        max_risk = risk_level

            # Update main table for this threat
            self.threat_cells[threat_name].config(text=max_risk)

    def calculate_likelihood_from_saved_data(self, threat_name, asset_key, asset_data):
        """Calculates likelihood from saved data using quadratic mean"""
        try:
            # Check if necessary values are present (Vulnerability, Access, Defense)
            if not all(str(i) in asset_data for i in [0, 1, 2]):
                return 0.0
            
            values = []
            for i in [0, 1, 2]:
                val = asset_data[str(i)]
                if not val or val == "0":
                    return 0.0
                values.append(float(val))
            
            if len(values) == 3:
                # Calculate likelihood using quadratic mean
                quadratic_mean = math.sqrt(sum(x**2 for x in values) / 3)
                likelihood = (quadratic_mean - 1) / 4  # Normalize [1,5] -> [0,1]
                return max(0.0, min(1.0, likelihood))
        
        except (ValueError, KeyError, TypeError):
            pass
        
        return 0.0
    def calculate_impact_from_saved_data(self, asset_data):
        """Calculates impact from saved data using quadratic mean"""
        try:
            # Check if necessary values are present (Operational Impact, Recovery)
            if not all(str(i) in asset_data for i in [3, 4]):
                return 0.0
            
            values = []
            for i in [3, 4]:  # Operational Impact, Recovery
                val = asset_data[str(i)]
                if not val or val == "0":
                    return 0.0
                values.append(float(val))
            
            if len(values) == 2:
                # Quadratic mean normalized
                quadratic_mean = math.sqrt(sum(x**2 for x in values) / 2)
                impact = (quadratic_mean - 1) / 4  # [1,5] -> [0,1]
                return max(0.0, min(1.0, impact))
        
        except (ValueError, KeyError, TypeError):
            pass
        
        return 0.0
    
    def update_main_table_risk_realtime(self):
        """Updates main table in real-time during calculations"""
        current_threat = self.selected_threat_var.get()
        if not current_threat or current_threat not in self.threat_cells:
            return
        
        # Find maximum risk among all currently displayed assets
        risk_priorities = {"Very High": 5, "High": 4, "Medium": 3, "Low": 2, "Very Low": 1, "": 0}
        max_risk = ""
        max_priority = 0
        
        # Check all currently displayed assets
        for key in self.impact_entries:
            risk_widget = self.impact_entries[key].get(7)  # Risk column
            if risk_widget and hasattr(risk_widget, 'cget'):
                risk_text = risk_widget.cget('text')
                priority = risk_priorities.get(risk_text, 0)
                if priority > max_priority:
                    max_priority = priority
                    max_risk = risk_text

        # Update main table in real time
        if max_risk:
            self.threat_cells[current_threat].config(text=max_risk)
    
    def get_max_risk_for_threat(self, threat_name):
        """Calculates the maximum risk for a specific threat"""
        if threat_name not in self.threat_data:
            return ""
        
        risk_priorities = {"Very High": 5, "High": 4, "Medium": 3, "Low": 2, "Very Low": 1, "": 0}
        max_risk = ""
        max_priority = 0

        # Calculate maximum risk for this threat
        threat_data = self.threat_data[threat_name]
        
        for asset_key, asset_data in threat_data.items():
            # Calculate likelihood for this asset
            likelihood = self.calculate_likelihood_from_saved_data(threat_name, asset_key, asset_data)

            # Calculate impact for this asset
            impact = self.calculate_impact_from_saved_data(asset_data)

            # Calculate risk if both are available
            if likelihood >= 0 and impact >= 0:
                likelihood_cat = self.value_to_category(likelihood)
                impact_cat = self.value_to_category(impact)
                risk_level = self.RISK_MATRIX.get((likelihood_cat, impact_cat), "")
                
                priority = risk_priorities.get(risk_level, 0)
                if priority > max_priority:
                    max_priority = priority
                    max_risk = risk_level
        
        return max_risk

    # ===== WORD EXPORT/IMPORT METHODS =====
    def export_to_word(self):
        """Exports the risk assessment to a Word document"""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Error", "python-docx library not available!\nInstall with: pip install python-docx")
            return
            
        try:
            # Mission type
            mission_type = self.mission_type_var.get()
            if mission_type == self.MISSION_TYPES[0]:  
                mission_type = ""  # Do not show if it is the default value

            # Automatic file name with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Risk_Assessment_0-A_{timestamp}.docx"

            # Create Output directory if it doesn't exist
            output_dir = os.path.join(get_base_path(), "Output")
            os.makedirs(output_dir, exist_ok=True)

            # Destination folder (Output directory)
            file_path = os.path.join(output_dir, filename)

            # Create Word document
            doc = Document()
            
            # Add content
            self.add_word_title(doc)
            if mission_type:
                self.add_mission_type(doc, mission_type)
            self.add_main_threats_table(doc)
            self.add_threat_details(doc)

            # Save document
            doc.save(file_path)
            
            messagebox.showinfo("Success", f"Risk Assessment exported to:\n{file_path}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error exporting to Word:\n{str(e)}")
    def import_from_word(self):
        """Import data from a previously exported Word document"""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Error", "python-docx library not available!\nInstall with: pip install python-docx")
            return
            
        try:
            file_path = filedialog.askopenfilename(
                filetypes=[("Word documents", "*.docx")],
                title="Import Risk Assessment from Word"
            )
            
            if not file_path:
                return            # Load Word document
            doc = Document(file_path)

            # Extract data
            self.extract_mission_type_from_word(doc)
            self.extract_threats_data_from_word(doc)

            # Update interface
            self.update_all_threats_in_main_table()
            
            messagebox.showinfo("Success", f"Risk Assessment imported from:\n{file_path}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error importing from Word:\n{str(e)}")
    
    def import_legacy_report(self):
        """Import data from a legacy Word report"""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Error", "python-docx library not available!\nInstall with: pip install python-docx")
            return
            
        try:
            file_path = filedialog.askopenfilename(
                filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
                title="Import Risk Assessment Report from older phases"
            )
            
            if not file_path:
                return
            
            # Load Word document
            doc = Document(file_path)
            
            # Parse the legacy report
            threats_data = self.parse_legacy_word_report(doc)
            
            if not threats_data:
                messagebox.showwarning("Warning", "No threat data found in the legacy report.")
                return
            
            # Clear existing data
            self.threat_data = {}
            
            # Import the parsed data into our data structure
            for threat_name, threat_info in threats_data.items():
                if threat_name in self.THREATS:
                    likelihood = threat_info['likelihood']
                    asset_categories = threat_info['assets']
                    
                    # Create threat data structure
                    threat_data = {}
                    
                    # For each asset category mentioned in the legacy report
                    for asset_category in asset_categories:
                        # Find the corresponding asset index in our ASSET_CATEGORIES
                        for i, (main_cat, sub_cat) in enumerate(self.ASSET_CATEGORIES):
                            if asset_category in sub_cat or asset_category in main_cat:
                                asset_key = f"{i+1}_probability"
                                
                                # Set minimal data to create a valid likelihood
                                # We'll use default values for Vulnerability, Access, Defense
                                # but set them to produce the desired likelihood
                                likelihood_val = self.category_to_value(likelihood)
                                base_val = self.likelihood_to_base_value(likelihood_val)
                                
                                # Use base_val - 1 for better calibration
                                adjusted_val = max(1, base_val)
                                if threat_name == "Jamming":
                                    threat_data[asset_key] = {
                                        '0': str(adjusted_val),  # Vulnerability
                                        '1': '',  # Access Control
                                        '2': str(adjusted_val),  # Defense Capability
                                        '3': '3',  # Operational Impact (medium)
                                        '4': '3'   # Recovery Time (medium)
                                }
                                else:
                                    threat_data[asset_key] = {
                                        '0': str(adjusted_val),  # Vulnerability
                                        '1': str(adjusted_val),  # Access Control
                                        '2': str(adjusted_val),  # Defense Capability
                                        '3': '3',  # Operational Impact (medium)
                                        '4': '3'   # Recovery Time (medium)
                                    }
                    
                    if threat_data:
                        self.threat_data[threat_name] = threat_data
            
            # Update interface
            self.update_all_threats_in_main_table()
            
            imported_count = len(threats_data)
            messagebox.showinfo("Success", f"Legacy report imported successfully!\n"
                                         f"Imported {imported_count} threats from:\n{file_path}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error importing legacy report:\n{str(e)}")
    
    def parse_legacy_report(self, content):
        """Parse a legacy text report and extract threat probability data"""
        threats_data = {}
        
        try:
            # Look for the "Main Threats Overview" section
            lines = content.split('\n')
            in_main_threats_section = False
            
            for line in lines:
                line = line.strip()
                
                # Check if we're entering the Main Threats Overview section
                if 'Main Threats Overview' in line:
                    in_main_threats_section = True
                    continue
                
                # Check if we're leaving the main threats section (entering detailed analysis)
                if in_main_threats_section and ('Detailed Threat Analysis' in line or 'Risk Matrix' in line):
                    break
                
                # If we're in the main threats section, look for threat data
                if in_main_threats_section and '|' in line:
                    # Skip header lines
                    if 'Threat' in line and 'Probability' in line:
                        continue
                    if line.startswith('|--') or line.startswith('---'):
                        continue
                    
                    # Parse threat line: "| Threat Name | Probability |"
                    parts = [part.strip() for part in line.split('|') if part.strip()]
                    
                    if len(parts) >= 2:
                        threat_name = parts[0].strip()
                        probability = parts[1].strip()
                        
                        # Map probability values to our standard values
                        probability_mapping = {
                            'Very Low': 'Very Low',
                            'Low': 'Low',
                            'Medium': 'Medium',
                            'High': 'High',
                            'Very High': 'Very High',
                            'VL': 'Very Low',
                            'L': 'Low',
                            'M': 'Medium',
                            'H': 'High',
                            'VH': 'Very High'
                        }
                        
                        # Normalize probability
                        probability = probability_mapping.get(probability, probability)
                        
                        # Only add if it's a valid threat name from our list
                        if threat_name in self.THREATS and probability in ['Very Low', 'Low', 'Medium', 'High', 'Very High']:
                            threats_data[threat_name] = probability
            
        except Exception as e:
            print(f"Error parsing legacy report: {e}")
        
        return threats_data
    
    def add_word_title(self, doc):
        """Adds the title to the Word document"""
        # Main title
        title = doc.add_heading('Risk Assessment', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data
        date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Empty space

    def add_mission_type(self, doc, mission_type=None):
        """Adds the mission type to the document"""
        if mission_type:
            mission_para = doc.add_paragraph(f'Mission Type: {mission_type}')
            mission_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Empty space

    def add_main_threats_table(self, doc):
        """Adds the main threats table"""
        doc.add_heading('Main Threats Overview', level=1)

        # Create simple table: Threat | Risk Level (as in main window)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Threat'
        header_cells[1].text = 'Risk Level'

        # Format header
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
        # Add data for each threat (use the same calculations as the main table)
        for threat in self.THREATS:
            row_cells = table.add_row().cells
            row_cells[0].text = threat

            # Calculate maximum risk for this threat
            max_risk = self.get_max_risk_for_threat(threat)
            row_cells[1].text = max_risk if max_risk else ""
        
        doc.add_page_break()
    def add_threat_details(self, doc):
        """Adds details for each threat with risk"""
        doc.add_heading('Detailed Threat Analysis', level=1)
        
        threats_with_data = []

        # First find all threats that have data
        for threat in self.THREATS:
            if threat in self.threat_data and self.threat_data[threat]:
                max_risk = self.get_max_risk_for_threat(threat)
                if max_risk:  # Only if a risk has been calculated
                    threats_with_data.append(threat)
        
        if not threats_with_data:
            doc.add_paragraph("No threats with assessed risk data.")
            return        # Add details for each threat with data
        for threat in threats_with_data:
            # Threat title
            doc.add_heading(f'{threat}', level=2)

            # Add asset table for this threat
            self.add_threat_asset_table(doc, threat)

            # Add countermeasures table for this threat
            self.add_threat_countermeasures_table(doc, threat)
            doc.add_paragraph()  # Empty space between threats
        
        # Add reference tables at the end
        doc.add_page_break()
        self.add_criteria_reference_table(doc)
        self.add_risk_matrix_table(doc)

    def add_threat_asset_table(self, doc, threat_name):
        """Creates asset assessment table for a specific threat"""
        doc.add_heading(f'Asset Assessment for {threat_name}', level=3)
        
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'

        # Exact header for import
        header_cells = table.rows[0].cells
        headers = ['Asset', 'Vulnerability', 'Access Control', 'Defense Capability', 
                  'Operational Impact', 'Recovery Time', 'Likelihood', 'Impact', 'Risk Level']
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Add rows for each asset with data
        assets_added = 0
        if threat_name in self.threat_data:
            threat_data = self.threat_data[threat_name]
            
            for asset_key, asset_data in threat_data.items():
                # Extract asset index from key (e.g., "1_probability" -> 0)
                try:
                    asset_index = int(asset_key.split('_')[0]) - 1
                    if 0 <= asset_index < len(self.ASSET_CATEGORIES):
                        category, asset_name = self.ASSET_CATEGORIES[asset_index]

                        # Calculate likelihood and impact
                        likelihood = self.calculate_likelihood_from_saved_data(threat_name, asset_key, asset_data)
                        impact = self.calculate_impact_from_saved_data(asset_data)

                        # Only if we have valid data
                        if asset_data and (likelihood >= 0 or impact >= 0):
                            row_cells = table.add_row().cells

                            # Asset name (important: must exactly match ASSET_CATEGORIES)
                            row_cells[0].text = asset_name

                            # Criteria (columns 0-4 correspond to Vulnerability, Access, Defense, Operational Impact, Recovery)
                            criteria_keys = ['0', '1', '2', '3', '4']
                            for j, key in enumerate(criteria_keys):
                                if key in asset_data:
                                    score = asset_data[key]
                                    # Specific format for import
                                    row_cells[j + 1].text = f"Score {score}"
                                else:
                                    row_cells[j + 1].text = "N/A"

                            # Calculate likelihood and impact if possible
                            if likelihood >= 0:
                                likelihood_cat = self.value_to_category(likelihood)
                                row_cells[6].text = likelihood_cat
                            else:
                                row_cells[6].text = "N/A"
                            
                            if impact >= 0:
                                impact_cat = self.value_to_category(impact)
                                row_cells[7].text = impact_cat
                            else:
                                row_cells[7].text = "N/A"
                            
                            # Risk Level
                            if likelihood >= 0 and impact >= 0:
                                likelihood_cat = self.value_to_category(likelihood)
                                impact_cat = self.value_to_category(impact)
                                risk_level = self.RISK_MATRIX.get((likelihood_cat, impact_cat), "N/A")
                                row_cells[8].text = risk_level
                            else:
                                row_cells[8].text = "N/A"
                            
                            assets_added += 1
                            
                except (ValueError, IndexError) as e:
                    print(f"Error processing asset {asset_key}: {e}")
                    continue
        
        # If no assets have data, add a placeholder row
        if assets_added == 0:
            row_cells = table.add_row().cells
            row_cells[0].text = "No asset data available"
            for i in range(1, 9):
                row_cells[i].text = "N/A"
        
        print(f"Added {assets_added} assets to table for threat {threat_name}")
        doc.add_paragraph()  # Space after the table
    
    def add_threat_countermeasures_table(self, doc, threat_name):
        """Add a table with security check for a specific threat"""
        # Add subtitle for controls
        doc.add_heading(f'Security Controls for {threat_name}', level=3)

        # Check if there are controls for this threat
        if threat_name not in self.THREAT_COUNTERMEASURES:
            doc.add_paragraph("No specific security controls defined for this threat.")
            return
        
        countermeasures = self.THREAT_COUNTERMEASURES[threat_name]

        # Create table with 2 columns (Control #, Control Name)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Control #'
        header_cells[1].text = 'Security Control'

        # Format header
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True

        # Add controls
        for i, control in enumerate(countermeasures, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = control

        # Add space after the table
        doc.add_paragraph()
    def extract_mission_type_from_word(self, doc):
        """Extracts the mission type from the Word document and updates the dropdown"""
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text.startswith("Mission Type:"):
                mission_type = text.replace("Mission Type:", "").strip()
                # Update the dropdown if the mission type is in the list
                if mission_type in self.MISSION_TYPES:
                    self.mission_type_var.set(mission_type)
                elif mission_type:  # If not in the list but not empty, use the first as fallback
                    self.mission_type_var.set(mission_type)
                    break
    def extract_threats_data_from_word(self, doc):
        """Extracts threat data from the Word document from the Detailed Threat Analysis section"""
        in_detailed_section = False

        # First, reset existing data to avoid conflicts
        self.threat_data = {}

        # Improved method: scan all document elements in order
        all_elements = []
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para_text = ""
                for para in doc.paragraphs:
                    if para._element == element:
                        para_text = para.text.strip()
                        break
                all_elements.append(('paragraph', para_text))
            elif element.tag.endswith('tbl'):  # Table
                for table in doc.tables:
                    if table._element == element:
                        all_elements.append(('table', table))
                        break

        # Now process elements in order
        current_threat = None
        threat_table_count = {}  # Count tables for each threat

        for element_type, element_data in all_elements:
            if element_type == 'paragraph':
                text = element_data

                # Check if we are in the "Detailed Threat Analysis" section
                if "Detailed Threat Analysis" in text:
                    in_detailed_section = True
                    print("[OK] Found Detailed Threat Analysis section")
                    continue

                # If we are in the detailed section, look for threat names
                if in_detailed_section and text in self.THREATS:
                    current_threat = text
                    threat_table_count[current_threat] = 0
                    print(f"[INFO] Found threat: {current_threat}")
                    
            elif element_type == 'table' and current_threat and in_detailed_section:
                table = element_data
                threat_table_count[current_threat] += 1
                table_number = threat_table_count[current_threat]
                
                print(f"[INFO] Processing table #{table_number} for threat: {current_threat}")

                # Check table type by number of columns
                if len(table.columns) == 9:
                    # Asset table
                    print(f"   -> Asset table detected (9 columns)")
                    self.extract_asset_table_data(table, current_threat)
                elif len(table.columns) == 2:
                    # Controls table (ignore for data import)
                    print(f"   -> Controls table detected (2 columns) - skipping")
                else:
                    print(f"   -> Unknown table format ({len(table.columns)} columns) - skipping")
                    
        print(f"[OK] Import completed. Found data for threats: {list(self.threat_data.keys())}")

        # Debug: show imported data
        for threat_name, threat_data in self.threat_data.items():
            print(f"   {threat_name}: {len(threat_data)} assets")
      
    def extract_asset_table_data(self, table, threat_name):
        """Extracts data from the asset table for a specific threat"""
        try:
            print(f"[INFO] Extracting asset table data for threat: {threat_name}")
            print(f"   Table dimensions: {len(table.rows)} rows x {len(table.columns)} columns")

            # Check table format (must have 9 columns)
            if len(table.columns) != 9:
                print(f"[ERROR] Invalid table format: expected 9 columns, got {len(table.columns)}")
                return

            # Check header to confirm it's the right table
            header_row = table.rows[0]
            expected_headers = ['Asset', 'Vulnerability', 'Access Control', 'Defense Capability', 
                              'Operational Impact', 'Recovery Time', 'Likelihood', 'Impact', 'Risk Level']
            
            header_match = True
            for i, expected in enumerate(expected_headers):
                if i < len(header_row.cells):
                    cell_text = header_row.cells[i].text.strip()
                    if expected.lower() not in cell_text.lower():
                        header_match = False
                        break
            
            if not header_match:
                print(f"[ERROR] Header mismatch - not an asset table")
                return
            
            print(f"[OK] Valid asset table confirmed")

            # Initialized data for this threat if not exists
            if threat_name not in self.threat_data:
                self.threat_data[threat_name] = {}

            # Process each data row (skip header)
            data_rows_processed = 0
            for row_idx, row in enumerate(table.rows[1:], 1):
                try:
                    cells = row.cells
                    if len(cells) < 9:
                        print(f"[ERROR]  Row {row_idx}: insufficient cells ({len(cells)})")
                        continue

                    # Extract asset name from the first cell
                    asset_name = cells[0].text.strip()
                    if not asset_name:
                        print(f"[ERROR]  Row {row_idx}: empty asset name")
                        continue
                    
                    print(f"   Processing asset: '{asset_name}'")

                    # Find the index of the corresponding asset in the standard categories
                    asset_index = None
                    for i, (category, name) in enumerate(self.ASSET_CATEGORIES):
                        if name.lower() == asset_name.lower():
                            asset_index = i + 1  # Index start from 1
                            break
                    
                    if asset_index is None:
                        print(f"[ERROR] Asset '{asset_name}' not found in standard categories")
                        # Try partial matching
                        for i, (category, name) in enumerate(self.ASSET_CATEGORIES):
                            if asset_name.lower() in name.lower() or name.lower() in asset_name.lower():
                                asset_index = i + 1
                                print(f"[OK] Found partial match: '{name}' -> using index {asset_index}")
                                break
                        
                        if asset_index is None:
                            continue

                    # Create key for threat data
                    asset_key = f"{asset_index}_probability"

                    # Extract scores from criteria (columns 1-5)
                    criteria_scores = {}
                    valid_scores = 0

                    for j in range(1, 6):  # Columns 1-5 for the 5 criteria
                        cell_text = cells[j].text.strip()

                        # Various parsing formats
                        score = self.parse_score_from_cell(cell_text)
                        
                        if score is not None:
                            criteria_scores[str(j-1)] = str(score)  # Save as string with index 0-4
                            valid_scores += 1
                            print(f"     Criterion {j-1}: {score}")
                        else:
                            print(f"     Criterion {j-1}: could not parse '{cell_text}'")

                    # Save only if we have at least 3 valid criteria (to calculate likelihood/impact)
                    if valid_scores >= 3:
                        self.threat_data[threat_name][asset_key] = criteria_scores
                        data_rows_processed += 1
                        print(f"[OK] Saved data for asset {asset_index} ({asset_name}): {valid_scores} criteria")
                    else:
                        print(f"[ERROR] Insufficient valid criteria ({valid_scores}/5) for asset {asset_name}")

                except Exception as e:
                    print(f"[ERROR] Error processing row {row_idx}: {e}")
                    continue
            
            print(f"[OK] Processed {data_rows_processed} valid asset rows for threat '{threat_name}'")
                    
        except Exception as e:
            print(f"[ERROR] Error extracting asset table data for {threat_name}: {e}")
    
    def parse_score_from_cell(self, cell_text):
        """Extracts a score from a Word table cell with various formats"""
        if not cell_text:
            return None
        
        text = cell_text.strip()

        # Format 1: "Score X"
        if "score" in text.lower():
            try:
                # Remove "Score" and take the number
                score_str = text.lower().replace("score", "").strip()
                return int(score_str)
            except ValueError:
                pass

        # Format 2: Only number
        if text.isdigit():
            score = int(text)
            if 1 <= score <= 5:  # Valida range
                return score

        # Format 3: Number in a longer string
        import re
        numbers = re.findall(r'\b([1-5])\b', text)
        if numbers:
            return int(numbers[0])

        # Format 4: "N/A" or empty
        if text.lower() in ['n/a', 'na', '-', '']:
            return None
        
        return None

    def add_criteria_reference_table(self, doc):
        """Adds the assessment criteria reference table"""
        doc.add_heading('Assessment Criteria Reference', level=1)

        # Creates criteria table (6 columns: Criteria + 5 score levels)
        table = doc.add_table(rows=len(self.CRITERIA_DATA), cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Populates the table with criteria data
        for i, row_data in enumerate(self.CRITERIA_DATA):
            row_cells = table.rows[i].cells
            for j, cell_text in enumerate(row_data):
                row_cells[j].text = cell_text

                # Formats the header (first row)
                if i == 0:
                    row_cells[j].paragraphs[0].runs[0].bold = True                # Formats the first column (criteria names)
                elif j == 0:
                    row_cells[j].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()  # Space after the table

    def add_risk_matrix_table(self, doc):
        """Adds the risk matrix ISO 27005"""
        doc.add_heading('Risk Assessment Matrix (ISO 27005)', level=1)

        # Defines the levels for the matrix (ISO 27005)
        levels = ["Very High", "High", "Medium", "Low", "Very Low"]

        # Creates 6x6 table (header + 5x5 matrix)
        table = doc.add_table(rows=6, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        # Empty cell in top left
        table.rows[0].cells[0].text = "Impact \n Likelihood"
        table.rows[0].cells[0].paragraphs[0].runs[0].bold = True

        # Header columns (Impact)
        for j, level in enumerate(levels, 1):
            table.rows[0].cells[j].text = level
            table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
        
        # Header rows (Likelihood) and matrix content
        for i, likelihood in enumerate(levels, 1):
            # Header row
            table.rows[i].cells[0].text = likelihood
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

            # Matrix content
            for j, impact in enumerate(levels, 1):
                risk_level = self.RISK_MATRIX.get((likelihood, impact), "")
                table.rows[i].cells[j].text = risk_level

                # Colors the cells based on risk level
                cell = table.rows[i].cells[j]
                if risk_level == "Very High":
                    # Dark Red
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(139, 0, 0)
                elif risk_level == "High":
                    # Red
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 20, 60)
                elif risk_level == "Medium":
                    # Orange
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 140, 0)
                elif risk_level == "Low":
                    # Dark Yellow
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(184, 134, 11)
                elif risk_level == "Very Low":
                    # Green
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(34, 139, 34)

        doc.add_paragraph()  # Space after the table

    def show_help(self):
        """Show help window with criteria descriptions"""
        help_window = tk.Toplevel(self.root)
        help_window.title("Assessment Criteria - Help")
        help_window.geometry("1200x700")  # Increased height to accommodate tool explanation
        help_window.configure(bg=self.COLORS['white'])
        help_window.resizable(True, True)
        
        # Center the window
        help_window.transient(self.root)
        help_window.grab_set()
        
        # Title
        title_label = tk.Label(help_window, text="Risk Assessment Criteria Descriptions", 
                              font=('Segoe UI',  16, 'bold'),
                              bg=self.COLORS['white'], fg=self.COLORS['dark'])
        title_label.pack(pady=(20, 15))
        
        # Create scrollable frame for the table
        canvas = tk.Canvas(help_window, bg=self.COLORS['white'], highlightthickness=0)
        scrollbar = tk.Scrollbar(help_window, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=self.COLORS['white'])
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Create table frame
        table_frame = tk.Frame(scrollable_frame, bg=self.COLORS['white'])
        table_frame.pack(fill='both', expand=True, padx=15, pady=10)
        
        # Table headers
        header_frame = tk.Frame(table_frame, bg=self.COLORS['primary'], relief='ridge', bd=1)
        header_frame.pack(fill='x', pady=(0, 2))
        
        # Configure grid for better column control
        header_frame.grid_columnconfigure(0, weight=0, minsize=250)
        header_frame.grid_columnconfigure(1, weight=1)
        
        criterion_header = tk.Label(header_frame, text="Criterion", font=('Segoe UI', 12, 'bold'),
                                   bg=self.COLORS['primary'], fg=self.COLORS['white'], anchor='w',
                                   padx=15, pady=10)
        criterion_header.grid(row=0, column=0, sticky='ew')
        
        desc_header = tk.Label(header_frame, text="Description", font=('Segoe UI', 12, 'bold'),
                              bg=self.COLORS['primary'], fg=self.COLORS['white'], anchor='w',
                              padx=15, pady=10)
        desc_header.grid(row=0, column=1, sticky='ew')
        
        # Criteria descriptions
        criteria_help = {
            "Vulnerability Level": "Measures the presence and severity of known security vulnerabilities in the system. Lower scores indicate well-patched systems with no known vulnerabilities, while higher scores indicate systems with actively exploitable vulnerabilities.",
            "Access Control": "Evaluates the strength of physical and logical access controls protecting the system. This includes authentication mechanisms, authorization policies, and physical security measures.",
            "Defense Capability": "Assesses the effectiveness of security countermeasures and detection systems. This includes firewalls, intrusion detection, monitoring systems, and incident response capabilities.",
            "Operational Impact": "Measures the potential impact on mission operations if the threat materializes. This considers service disruption, data loss, and effects on critical mission functions.",
            "Recovery Time": "Evaluates the time and resources required to restore normal operations after a security incident. This includes backup systems, recovery procedures, and business continuity planning."
        }
        
        # Add criteria rows
        for i, (criterion, description) in enumerate(criteria_help.items()):
            # Row frame
            row_color = self.COLORS['light'] if i % 2 == 0 else self.COLORS['white']
            row_frame = tk.Frame(table_frame, bg=row_color, relief='ridge', bd=1)
            row_frame.pack(fill='x', pady=1)
            
            # Configure grid for consistent column widths
            row_frame.grid_columnconfigure(0, weight=0, minsize=250)
            row_frame.grid_columnconfigure(1, weight=1)
            
            # Criterion name (left column)
            criterion_label = tk.Label(row_frame, text=criterion, 
                                      font=('Segoe UI', 11, 'bold'),
                                      bg=row_color, fg=self.COLORS['dark'], anchor='nw',
                                      padx=15, pady=8, wraplength=220, justify='left')
            criterion_label.grid(row=0, column=0, sticky='new')
            
            # Description (right column)
            desc_label = tk.Label(row_frame, text=description,
                                 font=('Segoe UI', 11),
                                 bg=row_color, fg='#495057', anchor='nw',
                                 padx=15, pady=8, wraplength=800, justify='left')
            desc_label.grid(row=0, column=1, sticky='new')
        
        # Add separator and tool explanation section
        separator_frame = tk.Frame(scrollable_frame, bg=self.COLORS['gray'], height=2)
        separator_frame.pack(fill='x', pady=(20, 15), padx=15)
        
        # Tool explanation title
        explanation_title = tk.Label(scrollable_frame, text="How the Risk Assessment Tool Works", 
                                    font=('Segoe UI', 14, 'bold'),
                                    bg=self.COLORS['white'], fg=self.COLORS['primary'])
        explanation_title.pack(pady=(10, 15), padx=15, anchor='w')
        
        # Tool explanation content
        explanation_frame = tk.Frame(scrollable_frame, bg=self.COLORS['light'], relief='ridge', bd=1)
        explanation_frame.pack(fill='x', padx=15, pady=(0, 20))
        
        explanation_text = """The Risk Assessment Tool for Phase 0-A helps evaluate cybersecurity risks during the preliminary design phase of space missions. Here's how to use it effectively:

1. MISSION CONFIGURATION:
   • Start by selecting your mission type from the dropdown (Earth Observation, Communication Satellite, etc.)
   • This helps contextualize the risk assessment for your specific mission profile

2. THREAT ANALYSIS PROCESS:
   • Click "THREAT ANALYSIS" to open the detailed assessment window
   • For each threat, evaluate all 9 asset categories (Ground Stations, Mission Control, Space Platform, etc.)
   • Rate each asset using the 5 assessment criteria on a scale of 1-5:
     - Vulnerability Level: How exposed is the asset to known vulnerabilities?
     - Access Control: How well protected is access to the asset?
     - Defense Capability: What countermeasures and detection systems are in place?
     - Operational Impact: How would a successful attack affect mission operations?
     - Recovery Time: How long would it take to restore normal operations?

3. AUTOMATIC RISK CALCULATION:
   • The tool automatically calculates Likelihood using the quadratic mean of Vulnerability, Access, and Defense scores
   • Impact is calculated using the quadratic mean of Operational Impact and Recovery Time scores
   • Likelihood is calculated using the quadratic mean of Vulnerability Level, Access Control and Defense Capability
   • Final Risk Level is determined using a standard risk matrix (Likelihood x Impact)
   • The main table shows the highest risk level found across all assets for each threat

4. DATA MANAGEMENT:
   • Save your assessments for each threat individually
   • Export complete reports to Word documents for documentation and sharing
   • Import previously saved assessments to continue work or review past analyses

5. LEGACY DATA INTEGRATION:
   • Use "IMPORT MISSION ANALYSIS REPORT" to import risk data from earlier mission phases
   • The tool can process reports from the Output folder containing mission analysis data
   • This helps maintain continuity across different project phases

6. OUTPUT AND REPORTING:
   • All exported reports are saved in the "Output" folder for easy access
   • Reports include detailed threat analysis, asset assessments, and security control recommendations
   • Check the Output folder for examples of risk reports from previous mission analysis phases

7. BEST PRACTICES:
   • Assess all relevant threats for your mission type
   • Consider both current security posture and planned security measures
   • Review and update assessments as the mission design evolves
   • Use the assessment results to prioritize security investments and controls"""
        
        explanation_label = tk.Label(explanation_frame, text=explanation_text,
                                   font=('Segoe UI', 10),
                                   bg=self.COLORS['light'], fg='#495057', anchor='nw',
                                   padx=20, pady=15, wraplength=1100, justify='left')
        explanation_label.pack(fill='both', expand=True)
        
        # Pack canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True, padx=(0, 5))
        scrollbar.pack(side="right", fill="y")
        
        # Setup global mouse wheel scrolling for the help window
        self.setup_global_mousewheel(scrollable_frame, canvas)
        
        # Focus on help window
        help_window.focus_set()
        
    def category_to_value(self, category):
        """Converts a category string to a numeric value (0-1 range)"""
        category_mapping = {
            'Very Low': 0.05,
            'Low': 0.1,
            'Medium': 0.4,
            'High': 0.7,
            'Very High': 0.9
        }
        return category_mapping.get(category, 0.5)
    
    def likelihood_to_base_value(self, likelihood_val):
        """Converts likelihood value to base assessment value (1-5 range)"""
        # Convert from [0,1] to [1,5] range
        # We need to reverse the formula: likelihood = (quadratic_mean - 1) / 4
        # So: quadratic_mean = likelihood * 4 + 1
        quadratic_mean = likelihood_val * 4 + 1
        
        # For simplicity, we'll use the quadratic mean as the base value
        # This assumes all three values (Vulnerability, Access, Defense) are equal
        base_value = max(1, min(5, int(round(quadratic_mean))))
        return base_value
    
    def parse_legacy_word_report(self, doc):
        """Parse a legacy Word report and extract threat data"""
        threats_data = {}
        
        try:
            # First, extract mission type from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text.startswith("Mission Type:"):
                    mission_type = text.replace("Mission Type:", "").strip()
                    # Update the dropdown if the mission type is in the list
                    if mission_type in self.MISSION_TYPES:
                        self.mission_type_var.set(mission_type)
                    elif mission_type:
                        # If not in the list, add it temporarily for this session
                        self.mission_type_var.set(mission_type)
            
            # Look for tables in the document
            for table in doc.tables:
                # Check if this is the main threats table
                if len(table.rows) > 1 and len(table.columns) >= 2:
                    # Check header row
                    header_row = table.rows[0]
                    headers = [cell.text.strip() for cell in header_row.cells]
                    
                    # Look for threat/probability table
                    if ('Threat' in headers and 'Probability' in headers) or \
                       ('Threat' in headers and any('Prob' in h for h in headers)):
                        
                        threat_col = -1
                        prob_col = -1
                        
                        # Find column indices
                        for i, header in enumerate(headers):
                            if 'Threat' in header:
                                threat_col = i
                            elif 'Probability' in header or 'Prob' in header:
                                prob_col = i
                        
                        if threat_col >= 0 and prob_col >= 0:
                            # Parse data rows
                            for row in table.rows[1:]:  # Skip header
                                if len(row.cells) > max(threat_col, prob_col):
                                    threat_name = row.cells[threat_col].text.strip()
                                    probability = row.cells[prob_col].text.strip()
                                    
                                    # Map probability values
                                    prob_mapping = {
                                        'Very Low': 'Very Low', 'Low': 'Low', 'Medium': 'Medium',
                                        'High': 'High', 'Very High': 'Very High',
                                        'VL': 'Very Low', 'L': 'Low', 'M': 'Medium',
                                        'H': 'High', 'VH': 'Very High'
                                    }
                                    
                                    probability = prob_mapping.get(probability, probability)
                                    
                                    # Only add if it's a valid threat
                                    if threat_name in self.THREATS and probability in ['Very Low', 'Low', 'Medium', 'High', 'Very High']:
                                        # Extract asset categories from detailed analysis
                                        asset_categories = self.extract_asset_categories_from_doc(doc, threat_name)
                                        
                                        threats_data[threat_name] = {
                                            'likelihood': probability,
                                            'assets': asset_categories
                                        }
            
        except Exception as e:
            print(f"Error parsing legacy Word report: {e}")
        
        return threats_data
    
    def extract_asset_categories_from_doc(self, doc, threat_name):
        """Extract asset categories mentioned for a specific threat"""
        asset_categories = []
        
        try:
            # Look for the threat in the document text
            in_threat_section = False
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                # Check if we're in the threat section
                if text == threat_name:
                    in_threat_section = True
                    continue
                elif in_threat_section and text in self.THREATS:
                    # We've moved to another threat section
                    break
                elif in_threat_section and text.startswith('Security Controls'):
                    # End of threat section
                    break
                
                # If we're in the threat section, look for asset mentions
                if in_threat_section and text:
                    # Look for asset categories in the text
                    for main_cat, sub_cat in self.ASSET_CATEGORIES:
                        if main_cat.lower() in text.lower() or sub_cat.lower() in text.lower():
                            if main_cat not in asset_categories:
                                asset_categories.append(main_cat)
                            if sub_cat not in asset_categories:
                                asset_categories.append(sub_cat)
        
        except Exception as e:
            print(f"Error extracting asset categories for {threat_name}: {e}")
        
        # If no specific assets found, default to common ones
        if not asset_categories:
            asset_categories = ['Ground', 'Space', 'Link', 'User']
        
        return asset_categories

def main():
    """Main function"""
    root = tk.Tk()
    app = RiskAssessmentTool(root)
    root.mainloop()

if __name__ == "__main__":
    main()
